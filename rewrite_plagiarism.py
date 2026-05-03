# -*- coding: utf-8 -*-
"""
Script viết lại các đoạn bị đánh đạo văn trong luận văn Lê Văn Hải
Chiến lược: Paraphrase + thêm trích dẫn + góc nhìn ứng dụng thực tế
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
import copy

doc_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main.docx"
out_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_v2.docx"

doc = Document(doc_path)

# Bảng thay thế: (chuỗi nhận dạng gốc, nội dung viết lại gốc)
REPLACEMENTS = {

    # 1. Câu mở đầu về Cách mạng 4.0
    "Trong bối cảnh cuộc Cách mạng công nghiệp 4.0, chuyển đổi số đã trở thành yêu cầu tất yếu":
    "Quá trình chuyển đổi số trong giai đoạn Cách mạng công nghiệp 4.0 đang đặt ra yêu cầu cấp thiết đối với mọi doanh nghiệp, buộc các tổ chức phải thay đổi phương thức quản lý và vận hành để duy trì sức cạnh tranh. Đặc biệt, sự phát triển mạnh mẽ của các công nghệ web hiện đại đã tạo cơ hội để các doanh nghiệp vừa và nhỏ tiếp cận hệ thống quản trị (ERP) với chi phí phù hợp, triển khai linh hoạt. Thay vì lưu trữ dữ liệu phân tán, phần mềm quản lý giúp tập trung hóa thông tin và tự động hóa các quy trình nghiệp vụ phức tạp, giảm đáng kể sai sót do yếu tố con người.",

    # 2. Định nghĩa HTML
    "HTML (Hypertext Markup Language – Ngôn ngữ đánh dấu siêu văn bản) là ngôn ngữ đánh dấu chuẩn dùng soạn thảo các tài liệu World Wide Web, chỉ rõ một trang Web được hiển thị như thế nào trong trình duyệt. Trong dự án này dùng HTML5.":
    "HTML5 là phiên bản mới nhất của ngôn ngữ đánh dấu siêu văn bản, đóng vai trò xây dựng cấu trúc nội dung cho các trang web hiện đại. Thay vì chỉ đơn thuần định dạng văn bản như các phiên bản trước, HTML5 bổ sung nhiều thẻ ngữ nghĩa như <article>, <section>, <nav> giúp tổ chức nội dung rõ ràng và hỗ trợ tốt cho SEO. Trong dự án ViralWindow, HTML5 được chọn vì khả năng tương thích đa trình duyệt và hỗ trợ đầy đủ các tính năng web app hiện đại [2].",

    # 3. Định nghĩa CSS
    "CSS (Cascading Style Sheets) là một ngôn ngữ quy định cách trình bày cho các tài liệu viết bằng HTML, XHTML, XML, SVG hay UML, v.v. CSS cung cấp nhiều thuộc tính trình bày cho các đối tượng với sự sáng tạo trong việc kết hợp các thuộc tính giúp mang lại hiệu quả cao.":
    "CSS3 là công nghệ tạo kiểu giao diện giúp tách biệt hoàn toàn phần trình bày ra khỏi cấu trúc HTML, mang lại sự linh hoạt và dễ bảo trì hơn cho dự án. Nhờ hệ thống Cascade và Specificity, CSS3 cho phép định nghĩa các quy tắc trình bày nhất quán trên toàn hệ thống. Trong ViralWindow, CSS3 được sử dụng để xây dựng bộ design system thống nhất cho tất cả các module từ dashboard, bảng dữ liệu đến form nhập liệu, đảm bảo giao diện nhất quán và dễ bảo trì [3].",

    # 4. Định nghĩa Responsive Web
    "Responsive web design (RWD) là một phương pháp thiết kế web nhằm tạo ra các trang web có khả năng thích ứng với mọi kích thước màn hình và thiết bị khác nhau. Điều này đảm bảo người dùng có trải nghiệm tốt nhất trên bất kỳ thiết bị nào từ máy tính để bàn đến điện thoại di động.":
    "Thiết kế giao diện thích ứng (Responsive Design) là phương pháp xây dựng giao diện web có khả năng tự điều chỉnh bố cục theo kích thước thiết bị của người dùng. Đối với hệ thống ViralWindow, do người dùng chủ yếu là nhân viên làm việc tại xưởng hoặc văn phòng với nhiều loại thiết bị khác nhau, việc áp dụng responsive design giúp đảm bảo trải nghiệm nhất quán trên máy tính để bàn, laptop và máy tính bảng mà không cần xây dựng ứng dụng riêng biệt cho từng nền tảng.",

    # 5. Định nghĩa Node.js
    "Node.js là một môi trường chạy (runtime) cho JavaScript phía máy chủ, được xây dựng trên engine V8 của Google Chrome. Không giống như JavaScript truyền thống chỉ chạy trên trình duyệt, Node.js cho phép lập trình viên sử dụng JavaScript để xây dựng các ứng dụng backend như web server, API, hệ thống realtime, và các dịch vụ mạng.":
    "Node.js là nền tảng runtime cho phép thực thi JavaScript ở phía máy chủ, được xây dựng trên V8 engine của Google Chrome. Điểm khác biệt cốt lõi của Node.js so với các nền tảng backend truyền thống là cơ chế xử lý bất đồng bộ (non-blocking I/O), giúp hệ thống phản hồi nhiều yêu cầu đồng thời mà không cần tạo thêm luồng xử lý mới. Đây là lý do nhóm lựa chọn Node.js cho hệ thống ViralWindow, khi cần xử lý đồng thời các kết nối từ nhiều nhân viên khác nhau và tích hợp thông báo thời gian thực qua Socket.IO [7].",

    # 6. Định nghĩa SQL
    "SQL là ngôn ngữ tiêu chuẩn dùng để làm việc với các hệ quản trị cơ sở dữ liệu quan hệ như MySQL, Oracle Database, và Microsoft SQL Server. SQL cho phép người dùng thực hiện các thao tác cơ bản như:":
    "SQL (Structured Query Language) là ngôn ngữ truy vấn chuẩn được dùng để tương tác với cơ sở dữ liệu quan hệ. Trong hệ thống ViralWindow, SQL được sử dụng để thực hiện đầy đủ các nghiệp vụ dữ liệu, bao gồm:",

    # 7. Định nghĩa MySQL
    "MySQL là một hệ quản trị cơ sở dữ liệu quan hệ (RDBMS – Relational Database Management System) mã nguồn mở, được sử dụng rộng rãi trong các ứng dụng web và hệ thống thông tin. MySQL sử dụng SQL (Structured Query Language) làm ngôn ngữ chính để truy vấn, quản lý và thao tác dữ liệu.":
    "MySQL 8.0 được lựa chọn làm hệ quản trị cơ sở dữ liệu chính cho dự án ViralWindow vì ba lý do cốt lõi: tính ổn định cao sau nhiều năm phát triển, khả năng tích hợp tốt với Node.js thông qua thư viện mysql2, và hỗ trợ đầy đủ các tính năng giao dịch (Transaction) cần thiết để đảm bảo tính nhất quán khi xử lý nghiệp vụ tài chính và kho vật tư. MySQL sử dụng SQL làm ngôn ngữ tương tác dữ liệu và hỗ trợ storage engine InnoDB đảm bảo tính toàn vẹn tham chiếu giữa các bảng [9].",

    # 8. Định nghĩa Express.js
    "Express.js là một framework web tối giản và linh hoạt được xây dựng trên nền tảng Node.js. Express.js cung cấp các công cụ và tính năng cần thiết để phát triển ứng dụng web và API một cách nhanh chóng, đơn giản và có tổ chức.":
    "Express.js được nhóm lựa chọn là framework backend vì tính tối giản, linh hoạt và phù hợp với kiến trúc RESTful API của hệ thống ViralWindow. Express.js không áp đặt cấu trúc cứng nhắc, cho phép nhóm tự tổ chức mã nguồn theo mô hình MVC phù hợp với yêu cầu dự án. Hệ thống middleware của Express.js cũng giúp tích hợp nhanh chóng các thành phần xác thực JWT, xử lý CORS và logging mà không cần viết lại từ đầu [8].",

    # 9. Mô tả Google Gemini AI
    "Google Generative AI là nền tảng trí tuệ nhân tạo do Google phát triển, cung cấp các mô hình AI có khả năng tạo sinh nội dung (Generative AI) như văn bản, hình ảnh, mã nguồn và nhiều dạng dữ liệu khác. Một trong những công nghệ cốt lõi của nền tảng này là Gemini – dòng mô hình AI đa phương thức (multimodal), có khả năng xử lý đồng thời nhiều loại dữ liệu như văn bản, hình ảnh, âm thanh và video.":
    "Gemini API là dịch vụ AI tổng quát (Generative AI) do Google phát triển, được tích hợp vào hệ thống ViralWindow nhằm xây dựng tính năng trợ lý ảo thông minh. Khác với các chatbot đơn giản, Gemini có khả năng hiểu ngữ cảnh nghiệp vụ cụ thể của ngành cửa nhôm kính và đưa ra gợi ý phân tích dữ liệu phù hợp. Trong phạm vi dự án, Gemini API được sử dụng để hỗ trợ nhân viên kinh doanh tra cứu thông tin kỹ thuật và phân tích báo cáo tài chính theo yêu cầu tự nhiên (natural language query).",
}

changed = 0
for para in doc.paragraphs:
    original = para.text.strip()
    for key, replacement in REPLACEMENTS.items():
        if key.strip() in original:
            # Giữ nguyên format, chỉ thay text
            if para.runs:
                # Xóa tất cả runs cũ, ghi vào run đầu tiên
                full_fmt = para.runs[0]
                for run in para.runs[1:]:
                    run.text = ""
                full_fmt.text = replacement
            else:
                para.add_run(replacement)
            changed += 1
            print(f"✅ Đã viết lại: {key[:60]}...")
            break

doc.save(out_path)
print(f"\n✅ Hoàn thành! Đã viết lại {changed} đoạn.")
print(f"📄 File mới: {out_path}")
