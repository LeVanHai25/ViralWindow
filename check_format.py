# -*- coding: utf-8 -*-
"""Kiểm tra toàn diện định dạng đồ án tốt nghiệp - ĐH Công Nghiệp Hà Nội"""
import sys, json
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_v2.docx"
doc = Document(doc_path)

# =============================================
# TIÊU CHUẨN ĐH CÔNG NGHIỆP HÀ NỘI
# =============================================
STD = {
    "margin_top_cm":    3.0,
    "margin_bottom_cm": 3.0,
    "margin_left_cm":   3.5,
    "margin_right_cm":  2.0,
    "font_body":        "Times New Roman",
    "font_size_body":   13,
    "line_spacing":     1.5,   # multiple
    "first_line_indent_cm": 1.25,
    "heading1_bold":    True,
    "heading1_upper":   True,
    "heading1_size":    13,
    "heading2_bold":    True,
    "heading2_size":    13,
    "align_body":       WD_ALIGN_PARAGRAPH.JUSTIFY,
}

issues = []
ok = []
warnings = []

# ------ 1. Kiểm tra lề trang ------
section = doc.sections[0]
top    = section.top_margin.cm    if section.top_margin    else None
bot    = section.bottom_margin.cm if section.bottom_margin else None
left   = section.left_margin.cm   if section.left_margin   else None
right  = section.right_margin.cm  if section.right_margin  else None

def check_margin(name, val, expected, tol=0.2):
    if val is None:
        issues.append(f"Lề {name}: không xác định được")
        return
    diff = abs(val - expected)
    if diff <= tol:
        ok.append(f"Lề {name}: {val:.1f}cm ✅ (chuẩn {expected}cm)")
    else:
        issues.append(f"Lề {name}: {val:.1f}cm ❌ (chuẩn {expected}cm, lệch {diff:.1f}cm)")

check_margin("trên",   top,   STD["margin_top_cm"])
check_margin("dưới",   bot,   STD["margin_bottom_cm"])
check_margin("trái",   left,  STD["margin_left_cm"])
check_margin("phải",   right, STD["margin_right_cm"])

# ------ 2. Kiểm tra từng đoạn văn ------
body_issues_font    = []
body_issues_size    = []
body_issues_spacing = []
body_issues_indent  = []
body_issues_align   = []
heading_issues      = []

para_count = 0
body_para_count = 0

HEADING_KEYWORDS = ["CHƯƠNG", "MỞ ĐẦU", "KẾT LUẬN", "TÀI LIỆU THAM KHẢO",
                    "DANH MỤC", "MỤC LỤC", "PHỤ LỤC"]

def is_heading(para):
    style_name = para.style.name.lower()
    txt = para.text.strip().upper()
    return ("heading" in style_name or
            any(kw in txt for kw in HEADING_KEYWORDS) or
            style_name.startswith("heading"))

def get_run_font(para):
    for run in para.runs:
        if run.font.name:
            return run.font.name
    if para.style and para.style.font and para.style.font.name:
        return para.style.font.name
    return None

def get_run_size(para):
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt
    if para.style and para.style.font and para.style.font.size:
        return para.style.font.size.pt
    return None

def get_line_spacing(para):
    pf = para.paragraph_format
    if pf.line_spacing:
        from docx.shared import Length
        ls = pf.line_spacing
        try:
            ls_val = ls / 12700 / 12  # convert EMU to points then to multiple of 12pt
            return round(ls_val, 2)
        except:
            return None
    return None

for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if not txt or len(txt) < 10:
        continue
    para_count += 1

    if is_heading(para):
        # Kiểm tra heading
        font_name = get_run_font(para)
        font_size = get_run_size(para)
        is_bold = any(run.bold for run in para.runs if run.text.strip())

        if font_name and "Times New Roman" not in font_name:
            heading_issues.append(f"Dòng {i+1}: Heading font '{font_name}' (cần Times New Roman)")
        if font_size and abs(font_size - 13) > 1:
            heading_issues.append(f"Dòng {i+1}: Heading cỡ {font_size}pt (cần 13pt) – '{txt[:50]}'")
        if not is_bold and len(txt) > 5:
            heading_issues.append(f"Dòng {i+1}: Heading chưa in đậm – '{txt[:50]}'")
        continue

    # Đoạn thân văn
    body_para_count += 1

    # Font
    fn = get_run_font(para)
    if fn and "Times New Roman" not in fn:
        body_issues_font.append(f"Dòng {i+1}: font '{fn}' – '{txt[:60]}'")

    # Cỡ chữ
    fs = get_run_size(para)
    if fs and abs(fs - 13) > 0.5:
        body_issues_size.append(f"Dòng {i+1}: {fs}pt – '{txt[:60]}'")

    # Căn lề
    align = para.paragraph_format.alignment
    if align is not None and align != WD_ALIGN_PARAGRAPH.JUSTIFY and align != WD_ALIGN_PARAGRAPH.LEFT:
        body_issues_align.append(f"Dòng {i+1}: align={align} – '{txt[:60]}'")

    # Thụt đầu dòng
    pf = para.paragraph_format
    fi = pf.first_line_indent
    if fi is not None:
        fi_cm = fi / 360000  # EMU to cm
        if abs(fi_cm - STD["first_line_indent_cm"]) > 0.3 and fi_cm > 0:
            body_issues_indent.append(f"Dòng {i+1}: indent={fi_cm:.2f}cm (cần 1.25cm) – '{txt[:50]}'")

# ------ 3. Kiểm tra đánh số thứ tự ------
fig_nums   = []
table_nums = []
for para in doc.paragraphs:
    t = para.text.strip()
    if t.lower().startswith("hình "):
        fig_nums.append(t[:60])
    if t.lower().startswith("bảng "):
        table_nums.append(t[:60])

# ------ 4. Kiểm tra tài liệu tham khảo ------
refs = []
in_refs = False
for para in doc.paragraphs:
    t = para.text.strip()
    if "TÀI LIỆU THAM KHẢO" in t.upper():
        in_refs = True
        continue
    if in_refs and t.startswith("["):
        refs.append(t[:80])

# ------ Tổng hợp ------
summary = {
    "total_paragraphs": para_count,
    "body_paragraphs": body_para_count,
    "margins_ok": [o for o in ok if "Lề" in o],
    "margins_issues": [o for o in issues if "Lề" in o],
    "font_issues_count": len(body_issues_font),
    "font_issues_sample": body_issues_font[:5],
    "size_issues_count": len(body_issues_size),
    "size_issues_sample": body_issues_size[:5],
    "align_issues_count": len(body_issues_align),
    "indent_issues_count": len(body_issues_indent),
    "indent_issues_sample": body_issues_indent[:5],
    "heading_issues_count": len(heading_issues),
    "heading_issues_sample": heading_issues[:5],
    "figure_count": len(fig_nums),
    "table_count": len(table_nums),
    "ref_count": len(refs),
    "refs_sample": refs[:5],
}

with open(r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\format_check.json","w",encoding="utf-8") as f:
    json.dump(summary, f, ensure_ascii=False, indent=2)

print(json.dumps(summary, ensure_ascii=False, indent=2))
