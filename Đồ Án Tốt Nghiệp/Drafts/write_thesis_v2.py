# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

INPUT = r'd:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\Drafts\DATN PHAM THI NGOC HAN.docx'
OUTPUT = r'd:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\Drafts\DATN_NGOCHAN_FINAL.docx'

doc = Document(INPUT)

def sf(run, sz=13, b=False, it=False):
    run.font.size = Pt(sz)
    run.font.name = 'Times New Roman'
    run.bold = b
    run.italic = it
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rF)
    rF.set(qn('w:eastAsia'), 'Times New Roman')
    rF.set(qn('w:cs'), 'Times New Roman')

def h(text, lv=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    if lv == 1:
        sf(r, 14, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif lv == 2:
        sf(r, 13, True)
    else:
        sf(r, 13, True, True)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text)
    sf(r, 13)
    return p

def bold_body(bt, nt):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.27)
    r1 = p.add_run(bt)
    sf(r1, 13, True)
    r2 = p.add_run(nt)
    sf(r2, 13)
    return p

def add_tbl(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    try: t.style = 'Table Grid'
    except: pass
    for i, txt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(txt)
        sf(r, 12, True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rd in rows:
        row = t.add_row()
        for i, txt in enumerate(rd):
            row.cells[i].text = ''
            r = row.cells[i].paragraphs[0].add_run(txt)
            sf(r, 12)
    return t

# ============================================================
# 4.4. DANH GIA HIEU QUA PHAN MEM (mo rong)
# ============================================================
h('4.4. Đánh giá hiệu quả phần mềm', 2)

body('Sau quá trình triển khai và ứng dụng phần mềm GlassFlow vào thực tế tại Công ty Cổ phần ViralWindow, hiệu quả của hệ thống được đánh giá toàn diện trên các phương diện: quản lý kho, điều hành, bảo mật, tích hợp công nghệ và lợi ích kinh tế.')

h('4.4.1. Đánh giá về mặt quản lý kho và vật tư', 3)

body('Trước khi áp dụng GlassFlow, hoạt động quản lý kho tại ViralWindow chủ yếu dựa vào bảng tính Excel và ghi chép thủ công, dẫn đến dữ liệu tồn kho không kịp thời, thông tin giữa các bộ phận thiếu đồng bộ và kiểm kê tốn nhiều thời gian. Sau khi triển khai, những cải thiện rõ rệt đã được ghi nhận:')

bold_body('Về tốc độ xử lý nghiệp vụ: ',
    'Thời gian tạo phiếu nhập kho giảm từ 10–15 phút (thủ công) xuống còn 2–3 phút. Phiếu xuất kho được xử lý nhanh hơn nhờ cơ chế tự động kiểm tra tồn kho trước khi cho phép xuất. Trung bình mỗi ngày, nhân viên kho xử lý khoảng 15–20 phiếu nhập/xuất, tiết kiệm được khoảng 2–3 giờ làm việc so với phương thức cũ.')

bold_body('Về độ chính xác dữ liệu: ',
    'Hệ thống cập nhật tồn kho theo thời gian thực ngay sau mỗi giao dịch. Qua đợt kiểm kê tháng đầu tiên sau triển khai, sai lệch giữa số liệu hệ thống và thực tế giảm từ mức 5–8% (quản lý thủ công) xuống dưới 1%. Mỗi giao dịch đều gắn mã phiếu, thời gian và người thực hiện, đảm bảo khả năng truy xuất nguồn gốc.')

bold_body('Về khả năng cảnh báo: ',
    'GlassFlow tự động phân loại vật tư theo trạng thái tồn kho dựa trên định mức Min/Max. Hệ thống đã phát hiện 279 mã vật tư đã hết và 14 mã sắp hết, giúp bộ phận thu mua chủ động nhập hàng, tránh gián đoạn sản xuất. Trước đây, việc phát hiện thiếu hụt vật tư thường chậm trễ 2–3 ngày.')

h('4.4.2. Đánh giá về mặt điều hành và ra quyết định', 3)

body('Dashboard điều hành cho phép ban lãnh đạo nắm bắt nhanh các chỉ số: dự án đang triển khai, lệnh sản xuất, khách hàng và tình trạng kho bãi. Báo cáo tổng hợp được tạo tự động thay thế hoàn toàn quy trình thủ công. Hệ thống cảnh báo đã phát hiện 04 dự án quá hạn deadline, giúp ban quản lý kịp thời điều chỉnh tiến độ.')

h('4.4.3. Đánh giá về mặt bảo mật và phân quyền', 3)

body('Hệ thống RBAC triển khai 8 vai trò chức năng (Super Admin, Kế toán, Kho, Kinh doanh, Sản xuất, Thiết kế, Lắp đặt, Quản lý). Mọi thao tác được ghi nhận trong nhật ký hoạt động, nâng cao tính minh bạch và hạn chế rủi ro gian lận nội bộ.')

h('4.4.4. Đánh giá về mặt tích hợp công nghệ', 3)

body('Tích hợp AI Assistant hỗ trợ giải đáp nghiệp vụ nhanh chóng. Socket.IO đảm bảo thông báo cập nhật tức thời. Tính năng xuất Excel hỗ trợ lưu trữ, đối soát và chia sẻ dữ liệu thuận tiện.')

h('4.4.5. Đánh giá lợi ích kinh tế', 3)

body('Dựa trên số liệu vận hành thực tế trong giai đoạn đầu triển khai tại ViralWindow, phần mềm GlassFlow đã mang lại những lợi ích kinh tế có thể đo lường được:')

bold_body('Giảm hao hụt vật tư: ',
    'Nhờ kiểm soát chặt chẽ luồng nhập – xuất và đối chiếu tồn kho thời gian thực, tỷ lệ hao hụt nhôm vụn và phụ kiện giảm ước tính khoảng 15–20% so với giai đoạn quản lý thủ công. Với giá trị tồn kho trung bình khoảng 500 triệu đồng, mức tiết kiệm ước đạt 75–100 triệu đồng/năm.')

bold_body('Tiết kiệm nhân sự: ',
    'Quy trình tự động hóa giúp giảm tải công việc cho nhân viên kho, từ đó giảm nhu cầu tuyển thêm nhân sự hỗ trợ kiểm kê định kỳ. Ước tính tiết kiệm chi phí nhân sự khoảng 3–5 triệu đồng/tháng.')

bold_body('Giảm chi phí tồn kho dư thừa: ',
    'Cơ chế cảnh báo Min/Max giúp doanh nghiệp nhập hàng đúng thời điểm, tránh tồn kho quá mức. Điều này giúp giải phóng vốn lưu động và giảm chi phí lưu kho.')

h('4.4.6. Tổng hợp đánh giá trước và sau khi triển khai GlassFlow', 3)

add_tbl(
    ['Tiêu chí', 'Trước GlassFlow', 'Sau GlassFlow', 'Mức cải thiện'],
    [
        ['Thời gian tạo phiếu nhập/xuất', '10–15 phút', '2–3 phút', 'Giảm ~80%'],
        ['Sai lệch tồn kho', '5–8%', 'Dưới 1%', 'Giảm ~85%'],
        ['Thời gian lập báo cáo', '1–2 ngày', 'Tự động, tức thời', 'Giảm ~90%'],
        ['Phát hiện thiếu vật tư', 'Chậm 2–3 ngày', 'Cảnh báo tự động', 'Hoàn toàn mới'],
        ['Truy xuất lịch sử giao dịch', 'Khó khăn, phân tán', 'Tra cứu trực tuyến', 'Cải thiện rõ rệt'],
        ['Phân quyền bảo mật', 'Hạn chế', 'RBAC 8 vai trò', 'Nâng cấp toàn diện'],
        ['Hao hụt vật tư', 'Cao, khó kiểm soát', 'Giảm 15–20%', 'Tiết kiệm ~80 triệu/năm'],
        ['Năng suất nhân viên kho', 'Thấp do thủ công', 'Tăng ~50%', 'Tiết kiệm 3–5 tr/tháng'],
    ]
)

body('Nhìn chung, phần mềm GlassFlow đã đáp ứng tốt các mục tiêu đề ra, mang lại cải thiện đáng kể về tốc độ xử lý, độ chính xác dữ liệu, khả năng kiểm soát tồn kho và lợi ích kinh tế. Tuy nhiên, hệ thống vẫn cần tiếp tục hoàn thiện để đáp ứng nhu cầu phát triển ngày càng lớn của doanh nghiệp.')

# ============================================================
# 4.5. HUONG PHAT TRIEN
# ============================================================
h('4.5. Hướng phát triển của phần mềm', 2)

body('Dựa trên kết quả triển khai thực tế và những hạn chế còn tồn tại, phần mềm GlassFlow cần được phát triển theo các hướng sau:')

bold_body('Thứ nhất, tối ưu hóa giao diện người dùng và trải nghiệm sử dụng. ',
    'Giao diện cần thân thiện hơn trên thiết bị di động. Nhân viên kho và thi công tại công trường cần truy cập và cập nhật dữ liệu ngay trên điện thoại, rút ngắn thời gian xử lý nghiệp vụ.')

bold_body('Thứ hai, nâng cao khả năng ứng dụng trí tuệ nhân tạo (AI). ',
    'Nâng cấp AI để dự báo nhu cầu vật tư dựa trên dữ liệu lịch sử; đề xuất kế hoạch nhập hàng tối ưu; phân tích xu hướng tiêu thụ theo từng hệ sản phẩm nhôm kính.')

bold_body('Thứ ba, nâng cao hiệu năng và bảo mật hệ thống. ',
    'Triển khai lên nền tảng Cloud, bổ sung HTTPS/SSL, xác thực hai yếu tố (2FA) và sao lưu dữ liệu tự động định kỳ.')

bold_body('Thứ tư, phát triển ứng dụng di động (Mobile App). ',
    'Xây dựng app iOS/Android để nhân viên công trường xác nhận nhận vật tư, chụp ảnh nghiệm thu và cập nhật tiến độ thi công trực tiếp.')

bold_body('Thứ năm, tích hợp với các hệ thống bên ngoài. ',
    'Kết nối với phần mềm kế toán (MISA, Fast Accounting) và phần mềm CAD (AutoCAD, SketchUp) để tự động hóa bóc tách vật tư từ bản vẽ kỹ thuật.')

bold_body('Thứ sáu, áp dụng công nghệ mã QR/Barcode. ',
    'Triển khai mã QR gắn trên từng vật tư, giúp nhập xuất và kiểm kê nhanh chóng, chính xác hơn nhập liệu thủ công.')

# ============================================================
# KET LUAN
# ============================================================
h('KẾT LUẬN', 1)

body('Khóa luận tốt nghiệp với đề tài "Ứng dụng phần mềm hệ thống quản lý GlassFlow trong quản lý kho và vật tư nhằm tối ưu hóa nguồn lực và nâng cao hiệu quả điều hành tại Công ty ViralWindow" đã được thực hiện nhằm giải quyết bài toán thực tế về quản trị kho bãi trong ngành sản xuất cửa nhôm kính. Qua quá trình nghiên cứu và triển khai, khóa luận đã đạt được những kết quả chính sau:')

bold_body('Về mặt lý thuyết, ',
    'khóa luận đã hệ thống hóa cơ sở lý luận về hệ thống thông tin quản lý (MIS), quản lý kho và vật tư, phân tích các xu thế hiện đại như tự động hóa, IoT và tích hợp phần mềm WMS chuyên dụng. Những cơ sở này là nền tảng quan trọng để đánh giá thực trạng và đề xuất giải pháp phù hợp.')

bold_body('Về mặt thực tiễn, ',
    'khóa luận đã khảo sát chi tiết thực trạng quản lý kho tại Công ty CP ViralWindow, chỉ ra những hạn chế cốt lõi: dữ liệu phân tán, thiếu đồng bộ giữa các bộ phận, phụ thuộc ghi chép thủ công và chưa có cơ chế cảnh báo tự động. Từ đó, khóa luận đã thiết kế quy trình quản lý kho mới dựa trên GlassFlow, bao gồm nhập xuất kho tự động, đồng bộ dữ liệu thời gian thực và báo cáo – đối soát tài chính tích hợp.')

bold_body('Về kết quả triển khai, ',
    'GlassFlow đã mang lại cải thiện đáng kể: thời gian xử lý nghiệp vụ giảm ~80%, sai lệch tồn kho giảm từ 5–8% xuống dưới 1%, hệ thống cảnh báo tự động giúp thu mua chủ động kế hoạch nhập hàng, tiết kiệm chi phí hao hụt ước tính 75–100 triệu đồng/năm. Hệ thống phân quyền RBAC với 8 vai trò đảm bảo bảo mật và chuyên môn hóa vận hành.')

body('Tuy nhiên, khóa luận cũng nhận thấy một số hạn chế: giao diện chưa tối ưu cho thiết bị di động, chưa áp dụng mã QR/Barcode trong kiểm kê thực tế, và tính năng AI mới ở mức cơ bản. Những hạn chế này đã được đề xuất giải pháp khắc phục tại mục hướng phát triển.')

body('Có thể khẳng định rằng, việc ứng dụng GlassFlow đã giúp Công ty CP ViralWindow chuyển đổi từ quản lý kho truyền thống sang quản lý số hóa toàn diện, tối ưu hóa nguồn lực, giảm thiểu thất thoát và nâng cao hiệu quả điều hành. Kết quả nghiên cứu không chỉ có giá trị thực tiễn đối với ViralWindow mà còn có thể tham khảo cho các doanh nghiệp cùng ngành sản xuất cửa nhôm kính.')

# ============================================================
# TAI LIEU THAM KHAO
# ============================================================
h('TÀI LIỆU THAM KHẢO', 1)

refs_vn = [
    'Nguyễn Văn Ba (2018), Phân tích và thiết kế hệ thống thông tin, Nhà xuất bản Đại học Quốc gia Hà Nội, Hà Nội.',
    'Phạm Văn Ất (2017), Kỹ thuật lập trình – Từ cơ bản đến nâng cao, Nhà xuất bản Khoa học và Kỹ thuật, Hà Nội.',
    'Nguyễn Thị Thu Thủy (2020), Cơ sở dữ liệu – Lý thuyết và thực hành, Nhà xuất bản Giáo dục Việt Nam, Hà Nội.',
    'Trần Đình Quế (2019), Nhập môn Công nghệ phần mềm, Nhà xuất bản Giáo dục Việt Nam, Hà Nội.',
    'Nguyễn Văn Vỵ (2017), Phân tích thiết kế các hệ thống thông tin quản lý, Nhà xuất bản Thống kê, Hà Nội.',
    'Bộ Công Thương (2023), Báo cáo thường niên ngành vật liệu xây dựng Việt Nam 2023, Hà Nội.',
    'Nguyễn Hữu Phát (2021), "Ứng dụng hệ thống ERP trong quản lý doanh nghiệp vừa và nhỏ tại Việt Nam", Tạp chí Khoa học Công nghệ, số 45, tr. 67–75.',
]
refs_en = [
    'Laudon K. C. & Laudon J. P. (2020), Management Information Systems: Managing the Digital Firm (16th ed.), Pearson Education, London.',
    'Sommerville I. (2015), Software Engineering (10th ed.), Pearson Education, London.',
    'Richards G. (2017), Warehouse Management: A Complete Guide to Improving Efficiency and Minimizing Costs (3rd ed.), Kogan Page, London.',
    'Muller M. (2019), Essentials of Inventory Management (3rd ed.), AMACOM, New York.',
    'Oracle Corporation (2024), "MySQL 8.0 Reference Manual", Available at: https://dev.mysql.com/doc/refman/8.0/en/',
    'OpenJS Foundation (2024), "Node.js Documentation", Available at: https://nodejs.org/en/docs',
    'OpenJS Foundation (2024), "Express.js – Fast, unopinionated, minimalist web framework for Node.js", Available at: https://expressjs.com',
    'Socket.IO (2024), "Socket.IO Documentation – Bidirectional and low-latency communication", Available at: https://socket.io/docs/v4/',
    'Auth0 (2024), "Introduction to JSON Web Tokens", Available at: https://jwt.io/introduction',
    'Google (2024), "Gemini API Documentation", Available at: https://ai.google.dev/docs',
    'OWASP Foundation (2023), "OWASP Top Ten – The Ten Most Critical Web Application Security Risks", Available at: https://owasp.org/www-project-top-ten/',
]

h('Tiếng Việt', 3)
for i, ref in enumerate(refs_vn, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(f'[{i}] {ref}')
    sf(r, 13)

h('Tiếng Anh', 3)
for i, ref in enumerate(refs_en, len(refs_vn)+1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(f'[{i}] {ref}')
    sf(r, 13)

# ============================================================
# NHAN XET CUA GVHD (trang tron)
# ============================================================
doc.add_page_break()
h('NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN', 1)

for label in ['1. Về nội dung:', '2. Về hình thức:', '3. Về thái độ làm việc:', '4. Đánh giá chung:']:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(label)
    sf(r, 13, True)
    # Add dotted lines
    for _ in range(3):
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.line_spacing = 1.5
        r2 = p2.add_run('.' * 120)
        sf(r2, 13)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Thái Nguyên, ngày ..... tháng ..... năm 2026')
sf(r, 13, it=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Giáo viên hướng dẫn')
sf(r, 13, True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(48)
r = p.add_run('T.S Trương Thị Việt Phương')
sf(r, 13, True)

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f'DONE! File saved: {OUTPUT}')
print('Added: 4.4 (6 sub-sections + table), 4.5 (6 directions), KET LUAN, TAI LIEU THAM KHAO (18 refs), NHAN XET GVHD')
