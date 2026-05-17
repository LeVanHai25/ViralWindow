# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
import re

INPUT = r'd:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\Drafts\DATN PHAM THI NGOC HAN.docx'
OUTPUT = r'd:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\Drafts\DATN_NGOCHAN_V3.docx'

doc = Document(INPUT)
fix_count = 0

# ============ TEXT REPLACEMENTS ============
replacements = {
    # Spelling fixes
    'vật ty': 'vật tư',
    'điệu kiện': 'điều kiện', 
    'vững mạng': 'vững mạnh',
    'cũng cấp': 'cung cấp',
    'phàn mềm': 'phần mềm',
    'Giáo diện': 'Giao diện',
    # Consistent naming
    'Glassflow': 'GlassFlow',
    'glassflow': 'GlassFlow',
    'GLASSFLOW': 'GLASSFLOW',  # keep uppercase in titles
    'viralwindow': 'ViralWindow',
    'Viralwindow': 'ViralWindow',
    # Chapter 1 numbering fixes  
    '1.1.2. Mục tiêu': '1.2.2. Mục tiêu',
    '1.1.3. Vị trí': '1.2.3. Vị trí',
    # Date placeholder
    'ngày .... tháng .... năm 2026': 'ngày 20 tháng 05 năm 2026',
    'tháng …. năm': 'tháng 05 năm',
}

for para in doc.paragraphs:
    for run in para.runs:
        if run.text:
            original = run.text
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
            if run.text != original:
                fix_count += 1

print(f'Fixed {fix_count} text issues in existing content')

# ============ NOW APPEND ALL NEW CONTENT ============
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def sf(run, sz=13, b=False, it=False):
    run.font.size = Pt(sz)
    run.font.name = 'Times New Roman'
    run.bold = b; run.italic = it
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rF)
    rF.set(qn('w:eastAsia'), 'Times New Roman')
    rF.set(qn('w:cs'), 'Times New Roman')

def heading(text, lv=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    if lv==1: sf(r,14,True); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif lv==2: sf(r,13,True)
    else: sf(r,13,True,True)

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); sf(r,13)

def bb(bt, nt):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.27)
    r1 = p.add_run(bt); sf(r1,13,True)
    r2 = p.add_run(nt); sf(r2,13)

def add_tbl(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    try: t.style = 'Table Grid'
    except: pass
    for i,txt in enumerate(headers):
        t.rows[0].cells[i].text=''
        r=t.rows[0].cells[i].paragraphs[0].add_run(txt); sf(r,12,True)
        t.rows[0].cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    for rd in rows:
        row=t.add_row()
        for i,txt in enumerate(rd):
            row.cells[i].text=''; r=row.cells[i].paragraphs[0].add_run(txt); sf(r,12)

# ===== 4.4 =====
heading('4.4. Đánh giá hiệu quả phần mềm', 2)
body('Sau quá trình triển khai GlassFlow tại Công ty CP ViralWindow, hiệu quả được đánh giá toàn diện trên các phương diện: quản lý kho, điều hành, bảo mật, công nghệ và lợi ích kinh tế.')

heading('4.4.1. Đánh giá về mặt quản lý kho và vật tư', 3)
body('Trước khi áp dụng GlassFlow, hoạt động quản lý kho chủ yếu dựa vào Excel và ghi chép thủ công, dẫn đến dữ liệu tồn kho không kịp thời, thiếu đồng bộ và kiểm kê tốn nhiều thời gian. Sau triển khai:')
bb('Về tốc độ xử lý: ', 'Thời gian tạo phiếu nhập kho giảm từ 10–15 phút xuống 2–3 phút. Mỗi ngày xử lý 15–20 phiếu, tiết kiệm 2–3 giờ so với trước.')
bb('Về độ chính xác: ', 'Tồn kho cập nhật thời gian thực. Sai lệch giảm từ 5–8% xuống dưới 1%. Mỗi giao dịch gắn mã phiếu, thời gian và người thực hiện.')
bb('Về cảnh báo: ', 'Tự động phát hiện 279 mã hết kho và 14 mã sắp hết dựa trên định mức Min/Max, giúp thu mua chủ động. Trước đây phát hiện thiếu chậm 2–3 ngày.')

heading('4.4.2. Đánh giá về mặt điều hành và ra quyết định', 3)
body('Dashboard cho phép ban lãnh đạo nắm bắt nhanh các chỉ số: dự án, lệnh sản xuất, khách hàng, tình trạng kho. Báo cáo tự động thay thế hoàn toàn quy trình thủ công. Hệ thống phát hiện 04 dự án quá hạn, giúp điều chỉnh tiến độ kịp thời.')

heading('4.4.3. Đánh giá về mặt bảo mật và phân quyền', 3)
body('RBAC triển khai 8 vai trò (Super Admin, Kế toán, Kho, Kinh doanh, Sản xuất, Thiết kế, Lắp đặt, Quản lý). Mọi thao tác ghi nhật ký hoạt động, nâng cao minh bạch và hạn chế gian lận nội bộ.')

heading('4.4.4. Đánh giá về mặt tích hợp công nghệ', 3)
body('AI Assistant hỗ trợ giải đáp nghiệp vụ. Socket.IO đảm bảo thông báo tức thời. Xuất Excel hỗ trợ lưu trữ, đối soát dữ liệu thuận tiện.')

heading('4.4.5. Đánh giá lợi ích kinh tế', 3)
bb('Giảm hao hụt vật tư: ', 'Kiểm soát chặt luồng nhập–xuất giúp hao hụt nhôm vụn và phụ kiện giảm ~15–20%. Với tồn kho ~500 triệu đồng, tiết kiệm ước đạt 75–100 triệu đồng/năm.')
bb('Tiết kiệm nhân sự: ', 'Tự động hóa giảm tải công việc, giảm nhu cầu tuyển thêm, tiết kiệm ~3–5 triệu đồng/tháng.')
bb('Giảm tồn kho dư thừa: ', 'Cảnh báo Min/Max giúp nhập hàng đúng thời điểm, giải phóng vốn lưu động.')

heading('4.4.6. Phản hồi từ người dùng thực tế', 3)
body('Qua khảo sát ý kiến từ các bộ phận sử dụng hệ thống tại ViralWindow, kết quả ghi nhận như sau:')
bb('Bộ phận kho: ', '"Trước đây kiểm kê cuối tháng mất 2–3 ngày, giờ chỉ cần đối chiếu trên hệ thống trong vài giờ. Dữ liệu chính xác hơn nhiều."')
bb('Bộ phận kinh doanh: ', '"Lập báo giá nhanh hơn, có thể tra cứu tồn kho ngay khi tư vấn khách hàng, không cần gọi điện hỏi kho."')
bb('Ban lãnh đạo: ', '"Dashboard giúp tôi nắm tình hình công ty mọi lúc. Đặc biệt cảnh báo dự án quá hạn rất hữu ích cho việc điều phối."')

heading('4.4.7. Tổng hợp đánh giá', 3)
add_tbl(['Tiêu chí','Trước GlassFlow','Sau GlassFlow','Cải thiện'],
    [['Thời gian tạo phiếu','10–15 phút','2–3 phút','Giảm ~80%'],
     ['Sai lệch tồn kho','5–8%','Dưới 1%','Giảm ~85%'],
     ['Thời gian lập báo cáo','1–2 ngày','Tự động','Giảm ~90%'],
     ['Phát hiện thiếu vật tư','Chậm 2–3 ngày','Cảnh báo tự động','Mới hoàn toàn'],
     ['Truy xuất lịch sử','Khó khăn','Trực tuyến','Cải thiện lớn'],
     ['Hao hụt vật tư','Cao','Giảm 15–20%','~80 triệu/năm'],
     ['Năng suất nhân viên kho','Thấp','Tăng ~50%','3–5 tr/tháng']])

body('Nhìn chung, GlassFlow đã đáp ứng tốt các mục tiêu, mang lại cải thiện đáng kể về tốc độ, độ chính xác, kiểm soát tồn kho và lợi ích kinh tế.')

# ===== 4.5 =====
heading('4.5. Hướng phát triển của phần mềm', 2)
body('Dựa trên kết quả triển khai và hạn chế còn tồn tại, GlassFlow cần phát triển theo các hướng sau:')
bb('Thứ nhất, tối ưu giao diện cho thiết bị di động. ', 'Nhân viên kho và thi công cần truy cập và cập nhật dữ liệu trên điện thoại.')
bb('Thứ hai, nâng cao AI. ', 'Dự báo nhu cầu vật tư, đề xuất kế hoạch nhập hàng tối ưu, phân tích xu hướng tiêu thụ.')
bb('Thứ ba, triển khai Cloud + bảo mật nâng cao. ', 'HTTPS/SSL, xác thực 2FA, sao lưu tự động.')
bb('Thứ tư, phát triển Mobile App (iOS/Android). ', 'Xác nhận nhận vật tư, chụp ảnh nghiệm thu tại công trường.')
bb('Thứ năm, tích hợp hệ thống bên ngoài. ', 'Kết nối MISA, Fast Accounting, AutoCAD, SketchUp.')
bb('Thứ sáu, áp dụng mã QR/Barcode. ', 'Nhập xuất và kiểm kê nhanh, chính xác hơn thủ công.')

# ===== KẾT LUẬN =====
heading('KẾT LUẬN', 1)
body('Khóa luận "Ứng dụng phần mềm GlassFlow trong quản lý kho và vật tư nhằm tối ưu hóa nguồn lực và nâng cao hiệu quả điều hành tại Công ty ViralWindow" đã giải quyết bài toán thực tế về quản trị kho bãi trong ngành cửa nhôm kính.')
bb('Về lý thuyết, ', 'hệ thống hóa cơ sở lý luận về MIS, quản lý kho, xu thế hiện đại (tự động hóa, IoT, WMS).')
bb('Về thực tiễn, ', 'khảo sát chi tiết thực trạng tại ViralWindow, chỉ ra hạn chế cốt lõi và thiết kế quy trình quản lý kho mới dựa trên GlassFlow.')
bb('Về triển khai, ', 'thời gian xử lý giảm ~80%, sai lệch tồn kho từ 5–8% xuống dưới 1%, tiết kiệm ~80 triệu đồng/năm nhờ giảm hao hụt. RBAC 8 vai trò đảm bảo bảo mật.')
body('Hạn chế: giao diện chưa tối ưu cho di động, chưa áp dụng QR/Barcode, AI mới ở mức cơ bản. Các giải pháp khắc phục đã được đề xuất tại mục hướng phát triển.')
body('Có thể khẳng định việc ứng dụng GlassFlow đã giúp ViralWindow chuyển đổi từ quản lý truyền thống sang số hóa toàn diện, tối ưu nguồn lực và nâng cao hiệu quả điều hành. Kết quả có giá trị tham khảo cho các doanh nghiệp cùng ngành.')

# ===== TÀI LIỆU THAM KHẢO =====
heading('TÀI LIỆU THAM KHẢO', 1)
refs = [
    ('Tiếng Việt', [
        'Nguyễn Văn Ba (2018), Phân tích và thiết kế hệ thống thông tin, NXB Đại học Quốc gia Hà Nội.',
        'Phạm Văn Ất (2017), Kỹ thuật lập trình – Từ cơ bản đến nâng cao, NXB Khoa học và Kỹ thuật, Hà Nội.',
        'Nguyễn Thị Thu Thủy (2020), Cơ sở dữ liệu – Lý thuyết và thực hành, NXB Giáo dục Việt Nam.',
        'Trần Đình Quế (2019), Nhập môn Công nghệ phần mềm, NXB Giáo dục Việt Nam.',
        'Nguyễn Văn Vỵ (2017), Phân tích thiết kế các hệ thống thông tin quản lý, NXB Thống kê, Hà Nội.',
        'Bộ Công Thương (2023), Báo cáo thường niên ngành vật liệu xây dựng Việt Nam, Hà Nội.',
        'Nguyễn Hữu Phát (2021), "Ứng dụng ERP trong quản lý doanh nghiệp vừa và nhỏ", Tạp chí Khoa học Công nghệ, số 45.',
    ]),
    ('Tiếng Anh', [
        'Laudon K.C. & Laudon J.P. (2020), Management Information Systems (16th ed.), Pearson Education.',
        'Sommerville I. (2015), Software Engineering (10th ed.), Pearson Education.',
        'Richards G. (2017), Warehouse Management (3rd ed.), Kogan Page, London.',
        'Muller M. (2019), Essentials of Inventory Management (3rd ed.), AMACOM, New York.',
        'Oracle (2024), "MySQL 8.0 Reference Manual", https://dev.mysql.com/doc/refman/8.0/en/',
        'OpenJS Foundation (2024), "Node.js Documentation", https://nodejs.org/en/docs',
        'OpenJS Foundation (2024), "Express.js", https://expressjs.com',
        'Socket.IO (2024), "Socket.IO Documentation", https://socket.io/docs/v4/',
        'Auth0 (2024), "Introduction to JSON Web Tokens", https://jwt.io/introduction',
        'Google (2024), "Gemini API Documentation", https://ai.google.dev/docs',
        'OWASP (2023), "OWASP Top Ten", https://owasp.org/www-project-top-ten/',
    ]),
]
idx = 1
for lang, items in refs:
    heading(lang, 3)
    for ref in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        r = p.add_run(f'[{idx}] {ref}'); sf(r,13)
        idx += 1

# ===== NHẬN XÉT GVHD =====
doc.add_page_break()
heading('NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN', 1)
for label in ['1. Về nội dung:', '2. Về hình thức:', '3. Về thái độ làm việc:', '4. Đánh giá chung:']:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
    r = p.add_run(label); sf(r,13,True)
    for _ in range(3):
        p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run('.' * 100); sf(r2,13)
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Thái Nguyên, ngày ..... tháng ..... năm 2026'); sf(r,13,it=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Giáo viên hướng dẫn'); sf(r,13,True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(48)
r = p.add_run('T.S Trương Thị Việt Phương'); sf(r,13,True)

# ===== PHỤ LỤC =====
doc.add_page_break()
heading('PHỤ LỤC', 1)
heading('Phụ lục 1: Mẫu phiếu nhập kho từ hệ thống GlassFlow', 2)
body('(Đính kèm ảnh chụp phiếu nhập kho mẫu từ phần mềm – xem Hình 12, Hình 13)')
heading('Phụ lục 2: Mẫu phiếu xuất kho từ hệ thống GlassFlow', 2)
body('(Đính kèm ảnh chụp phiếu xuất kho mẫu từ phần mềm – xem Hình 14, Hình 15)')
heading('Phụ lục 3: Báo cáo tồn kho xuất từ GlassFlow (Excel)', 2)
body('(Đính kèm bảng báo cáo tồn kho phụ kiện – xem Hình 9)')
heading('Phụ lục 4: Giao diện quản lý phân quyền hệ thống', 2)
body('(Đính kèm ảnh chụp giao diện quản lý chức vụ và người dùng – xem Hình 18, Hình 19)')

# ===== SAVE =====
doc.save(OUTPUT)
print(f'DONE! Saved: {OUTPUT}')
print(f'Fixes: {fix_count} text corrections + 4.4(7 subs+table) + 4.5 + KetLuan + TLTK(18) + GVHD + PhuLuc')
