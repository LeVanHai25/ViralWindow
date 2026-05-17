# -*- coding: utf-8 -*-
"""
Script chen noi dung muc 4.4, 4.5 va KET LUAN vao file DATN
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

INPUT_PATH = r'd:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\Drafts\DATN PHAM THI NGOC HAN.docx'
OUTPUT_PATH = r'd:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\Drafts\DATN PHAM THI NGOC HAN_COMPLETE.docx'

doc = Document(INPUT_PATH)

# ===================== HELPERS =====================
def set_run_font(run, size=13, bold=False, italic=False, font_name='Times New Roman'):
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def add_heading_paragraph(doc, text, level=2):
    """Add a heading-style paragraph (bold, larger font)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=14, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        set_run_font(run, size=13, bold=True)
    elif level == 3:
        set_run_font(run, size=13, bold=True, italic=True)
    return p

def add_body_paragraph(doc, text, indent_first=True):
    """Add a normal body paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_run_font(run, size=13)
    return p

def add_body_with_bold_prefix(doc, bold_text, normal_text):
    """Add paragraph with bold prefix then normal text"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.27)
    r1 = p.add_run(bold_text)
    set_run_font(r1, size=13, bold=True)
    r2 = p.add_run(normal_text)
    set_run_font(r2, size=13)
    return p

def add_bullet(doc, text, level=0):
    """Add a bullet point"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run('– ' + text)
    set_run_font(run, size=13)
    return p

def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_run_font(run, size=12, bold=bold)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return row


# ===================== FIND INSERTION POINT =====================
# Find the paragraph that contains "4.4." or "4.5." or "Hướng phát triển"
insert_index = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '4.4.' in text and 'Đánh giá' in text:
        insert_index = i
        break
    if '4.5.' in text and 'Hướng phát triển' in text:
        insert_index = i
        break

if insert_index is None:
    # Find last paragraph as fallback
    insert_index = len(doc.paragraphs) - 1
    print(f"Warning: Could not find 4.4/4.5, appending at end (index {insert_index})")
else:
    print(f"Found insertion point at paragraph index {insert_index}")

# We'll append content at the end of the document since inserting in the middle is complex
# First, let's check if content already exists after 4.4
print("\n--- Adding content to document ---\n")

# ===================== SECTION 4.4 =====================
add_heading_paragraph(doc, '4.4. Đánh giá hiệu quả phần mềm', level=2)

add_body_paragraph(doc, 
    'Sau quá trình triển khai và ứng dụng phần mềm GlassFlow vào thực tế tại Công ty Cổ phần ViralWindow, '
    'hiệu quả của hệ thống được đánh giá trên các phương diện sau:')

# 4.4.1
add_heading_paragraph(doc, '4.4.1. Đánh giá về mặt quản lý kho và vật tư', level=3)

add_body_paragraph(doc,
    'Trước khi áp dụng phần mềm GlassFlow, hoạt động quản lý kho tại ViralWindow chủ yếu dựa vào bảng tính Excel '
    'và ghi chép thủ công. Điều này dẫn đến nhiều hạn chế như: dữ liệu tồn kho không được cập nhật kịp thời, '
    'thông tin giữa các bộ phận thiếu đồng bộ và việc kiểm kê tốn nhiều thời gian. Sau khi triển khai GlassFlow, '
    'những cải thiện rõ rệt đã được ghi nhận:')

add_body_with_bold_prefix(doc, 'Về tốc độ xử lý nghiệp vụ: ',
    'Thời gian tạo phiếu nhập kho giảm từ 10–15 phút (thủ công) xuống còn 2–3 phút trên hệ thống. '
    'Phiếu xuất kho được xử lý nhanh hơn nhờ cơ chế tự động kiểm tra tồn kho trước khi cho phép xuất, '
    'hạn chế tình trạng xuất vượt số lượng cho phép.')

add_body_with_bold_prefix(doc, 'Về độ chính xác dữ liệu: ',
    'Hệ thống cập nhật tồn kho theo thời gian thực ngay sau mỗi giao dịch nhập hoặc xuất. '
    'Sai lệch giữa số liệu trên hệ thống và thực tế kiểm kê giảm đáng kể so với phương thức quản lý trước đây. '
    'Mỗi giao dịch đều được gắn mã phiếu, thời gian và người thực hiện, đảm bảo khả năng truy xuất nguồn gốc.')

add_body_with_bold_prefix(doc, 'Về khả năng cảnh báo: ',
    'GlassFlow tự động phân loại vật tư theo trạng thái tồn kho (đủ hàng, sắp hết, đã hết) '
    'dựa trên định mức tối thiểu và tối đa được thiết lập. Hệ thống đã phát hiện và cảnh báo 279 mã vật tư '
    'đã hết kho và 14 mã sắp hết, giúp bộ phận thu mua chủ động lên kế hoạch nhập hàng kịp thời, '
    'tránh gián đoạn sản xuất.')

# 4.4.2
add_heading_paragraph(doc, '4.4.2. Đánh giá về mặt điều hành và ra quyết định', level=3)

add_body_paragraph(doc,
    'Phần mềm GlassFlow cung cấp giao diện Dashboard điều hành trực quan, cho phép ban lãnh đạo nắm bắt nhanh '
    'các chỉ số quan trọng như: số lượng dự án đang triển khai, lệnh sản xuất hiện tại, số lượng khách hàng '
    'và tình trạng kho bãi. Các báo cáo tổng hợp về tồn kho, lịch sử giao dịch và giá trị hàng tồn kho '
    'được tạo tự động, thay thế hoàn toàn quy trình lập báo cáo thủ công trước đây.')

add_body_paragraph(doc,
    'Đặc biệt, hệ thống cảnh báo rủi ro đã phát hiện 04 dự án quá hạn deadline, giúp ban quản lý kịp thời '
    'điều chỉnh tiến độ và phân bổ nguồn lực phù hợp. Tính năng này cho thấy GlassFlow không chỉ đơn thuần là '
    'công cụ quản lý kho mà còn đóng vai trò hỗ trợ ra quyết định điều hành ở cấp quản lý.')

# 4.4.3
add_heading_paragraph(doc, '4.4.3. Đánh giá về mặt bảo mật và phân quyền', level=3)

add_body_paragraph(doc,
    'Hệ thống phân quyền theo mô hình RBAC (Role-Based Access Control) đã được triển khai hiệu quả với 8 vai trò '
    'chức năng (Super Admin, Kế toán, Kho, Kinh doanh, Sản xuất, Thiết kế, Lắp đặt, Quản lý). Mỗi nhân viên chỉ '
    'được truy cập vào các module phù hợp với nhiệm vụ được phân công, đảm bảo tính bảo mật dữ liệu và nguyên tắc '
    'đặc quyền tối thiểu (Least Privilege).')

add_body_paragraph(doc,
    'Ngoài ra, mọi thao tác trên hệ thống đều được ghi nhận trong nhật ký hoạt động, tạo điều kiện cho việc '
    'kiểm tra và truy vết khi cần thiết. Điều này góp phần nâng cao tính minh bạch trong quản lý và hạn chế '
    'rủi ro gian lận nội bộ.')

# 4.4.4
add_heading_paragraph(doc, '4.4.4. Đánh giá về mặt tích hợp công nghệ', level=3)

add_body_paragraph(doc,
    'Việc tích hợp trợ lý AI (AI Assistant) vào hệ thống là một điểm sáng nổi bật, giúp nhân viên được hỗ trợ '
    'giải đáp nghiệp vụ nhanh chóng mà không cần liên hệ trực tiếp với bộ phận kỹ thuật. Hệ thống giao tiếp '
    'thời gian thực (Socket.IO) đảm bảo thông báo và dữ liệu được cập nhật tức thời giữa các bộ phận. '
    'Tính năng xuất báo cáo Excel hỗ trợ doanh nghiệp lưu trữ, đối soát và chia sẻ dữ liệu một cách thuận tiện.')

# 4.4.5 - Bảng tổng hợp
add_heading_paragraph(doc, '4.4.5. Tổng hợp đánh giá', level=3)

# Create comparison table
table = doc.add_table(rows=1, cols=4)
try:
    table.style = 'Table Grid'
except KeyError:
    pass  # style not available, use default

# Header row
hdr = table.rows[0]
headers = ['Tiêu chí đánh giá', 'Trước GlassFlow', 'Sau GlassFlow', 'Mức cải thiện']
for i, h in enumerate(headers):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    run = p.add_run(h)
    set_run_font(run, size=12, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
data = [
    ['Thời gian tạo phiếu nhập/xuất', '10–15 phút', '2–3 phút', 'Giảm ~80%'],
    ['Độ chính xác tồn kho', 'Sai lệch thường xuyên', 'Cập nhật thời gian thực', 'Cải thiện rõ rệt'],
    ['Thời gian lập báo cáo', '1–2 ngày', 'Tự động, tức thời', 'Giảm ~90%'],
    ['Khả năng cảnh báo thiếu hàng', 'Không có', 'Tự động theo định mức', 'Hoàn toàn mới'],
    ['Khả năng truy xuất lịch sử', 'Khó khăn, phân tán', 'Tra cứu trực tuyến', 'Cải thiện rõ rệt'],
    ['Phân quyền bảo mật', 'Hạn chế', 'RBAC 8 vai trò', 'Nâng cấp toàn diện'],
]
for row_data in data:
    add_table_row(table, row_data)

add_body_paragraph(doc,
    'Nhìn chung, phần mềm GlassFlow đã đáp ứng tốt các mục tiêu đề ra trong giai đoạn đầu triển khai, '
    'mang lại những cải thiện đáng kể về tốc độ xử lý, độ chính xác dữ liệu và khả năng kiểm soát tồn kho. '
    'Tuy nhiên, hệ thống vẫn cần được tiếp tục hoàn thiện và mở rộng để đáp ứng tốt hơn nhu cầu phát triển '
    'ngày càng lớn của doanh nghiệp.')

# ===================== SECTION 4.5 =====================
add_heading_paragraph(doc, '4.5. Hướng phát triển của phần mềm', level=2)

add_body_paragraph(doc,
    'Dựa trên kết quả triển khai thực tế và những hạn chế còn tồn tại, phần mềm GlassFlow cần được phát triển '
    'theo các hướng sau nhằm nâng cao hơn nữa hiệu quả quản lý kho và vật tư tại Công ty CP ViralWindow:')

add_body_with_bold_prefix(doc, 'Thứ nhất, tối ưu hóa giao diện người dùng và trải nghiệm sử dụng. ',
    'Giao diện phần mềm cần được cải thiện để thân thiện hơn với người dùng trên nhiều loại thiết bị, '
    'đặc biệt là thiết bị di động. Nhân viên kho và nhân viên thi công tại công trường cần có khả năng truy cập '
    'hệ thống và cập nhật dữ liệu ngay trên điện thoại thông minh, giúp rút ngắn thời gian xử lý nghiệp vụ '
    'và đảm bảo tính kịp thời của thông tin.')

add_body_with_bold_prefix(doc, 'Thứ hai, nâng cao khả năng ứng dụng trí tuệ nhân tạo (AI). ',
    'Hiện tại, tính năng AI Assistant mới dừng ở mức hỗ trợ giải đáp nghiệp vụ cơ bản. Trong tương lai, '
    'hệ thống cần được nâng cấp để AI có thể: dự báo nhu cầu vật tư dựa trên dữ liệu lịch sử và xu hướng '
    'đơn hàng; đề xuất kế hoạch nhập hàng tối ưu nhằm giảm thiểu chi phí lưu kho; phân tích xu hướng tiêu thụ '
    'vật tư theo từng hệ sản phẩm nhôm kính để hỗ trợ ban lãnh đạo ra quyết định chiến lược.')

add_body_with_bold_prefix(doc, 'Thứ ba, nâng cao hiệu năng và bảo mật hệ thống. ',
    'Hệ thống cần được triển khai lên nền tảng điện toán đám mây (Cloud) để tăng khả năng mở rộng và đảm bảo '
    'tính sẵn sàng cao. Đồng thời, cần bổ sung các cơ chế bảo mật nâng cao như: giao thức HTTPS/SSL cho toàn bộ '
    'kết nối, xác thực hai yếu tố (2FA) cho tài khoản quản trị và cơ chế sao lưu dữ liệu tự động định kỳ.')

add_body_with_bold_prefix(doc, 'Thứ tư, phát triển ứng dụng di động (Mobile App). ',
    'Việc xây dựng ứng dụng di động trên nền tảng iOS và Android sẽ giúp nhân viên tại công trường có thể: '
    'xác nhận nhận vật tư tại hiện trường, chụp ảnh nghiệm thu và cập nhật tiến độ thi công trực tiếp trên '
    'điện thoại. Điều này đặc biệt cần thiết trong bối cảnh ViralWindow đang mở rộng hệ thống đại lý và '
    'chi nhánh trên nhiều tỉnh thành.')

add_body_with_bold_prefix(doc, 'Thứ năm, tích hợp với các hệ thống bên ngoài. ',
    'GlassFlow cần được kết nối với phần mềm kế toán phổ biến (MISA, Fast Accounting) để đồng bộ dữ liệu '
    'tài chính, giảm thiểu việc nhập liệu trùng lặp. Ngoài ra, việc tích hợp với phần mềm thiết kế CAD '
    '(AutoCAD, SketchUp) sẽ cho phép tự động hóa quy trình bóc tách vật tư từ bản vẽ kỹ thuật, nâng cao '
    'độ chính xác và tiết kiệm thời gian cho bộ phận kỹ thuật.')

add_body_with_bold_prefix(doc, 'Thứ sáu, áp dụng công nghệ mã QR/Barcode trong quản lý kho thực tế. ',
    'Việc triển khai hệ thống mã QR hoặc mã vạch gắn trên từng vật tư sẽ giúp nhân viên kho thực hiện '
    'thao tác nhập xuất và kiểm kê nhanh chóng, chính xác hơn so với nhập liệu thủ công, đồng thời giảm thiểu '
    'sai sót trong quá trình quản lý danh mục vật tư đa dạng của ngành nhôm kính.')

# ===================== KẾT LUẬN =====================
add_heading_paragraph(doc, 'KẾT LUẬN', level=1)

add_body_paragraph(doc,
    'Khóa luận tốt nghiệp với đề tài "Ứng dụng phần mềm hệ thống quản lý GlassFlow trong quản lý kho và vật tư '
    'nhằm tối ưu hóa nguồn lực và nâng cao hiệu quả điều hành tại Công ty ViralWindow" đã được thực hiện nhằm '
    'giải quyết bài toán thực tế về quản trị kho bãi trong ngành sản xuất cửa nhôm kính. Qua quá trình nghiên cứu '
    'và triển khai, khóa luận đã đạt được những kết quả chính sau:')

add_body_with_bold_prefix(doc, 'Về mặt lý thuyết, ',
    'khóa luận đã hệ thống hóa cơ sở lý luận về hệ thống thông tin quản lý (MIS), quản lý kho và vật tư, '
    'đồng thời phân tích các xu thế quản lý kho hiện đại như tự động hóa, ứng dụng IoT và tích hợp phần mềm '
    'WMS chuyên dụng. Những cơ sở lý thuyết này là nền tảng quan trọng để đánh giá thực trạng và đề xuất '
    'giải pháp phù hợp cho doanh nghiệp.')

add_body_with_bold_prefix(doc, 'Về mặt thực tiễn, ',
    'khóa luận đã khảo sát và phân tích chi tiết thực trạng quản lý kho tại Công ty CP ViralWindow, chỉ ra '
    'những hạn chế cốt lõi như: dữ liệu phân tán, thiếu đồng bộ giữa các bộ phận, phụ thuộc vào ghi chép '
    'thủ công và chưa có cơ chế cảnh báo tồn kho tự động. Từ đó, khóa luận đã thiết kế quy trình quản lý kho '
    'mới dựa trên việc ứng dụng phần mềm GlassFlow, bao gồm: quy trình nhập xuất kho tự động, hệ thống '
    'đồng bộ dữ liệu thời gian thực và cơ chế báo cáo – đối soát tài chính tích hợp.')

add_body_with_bold_prefix(doc, 'Về kết quả triển khai, ',
    'phần mềm GlassFlow đã được ứng dụng thực tế tại ViralWindow và mang lại những cải thiện đáng kể: '
    'thời gian xử lý nghiệp vụ kho giảm khoảng 80%, độ chính xác dữ liệu tồn kho được nâng cao nhờ cập nhật '
    'thời gian thực, hệ thống cảnh báo tự động giúp bộ phận thu mua chủ động trong kế hoạch nhập hàng. '
    'Hệ thống phân quyền RBAC với 8 vai trò chức năng đảm bảo tính bảo mật và chuyên môn hóa trong vận hành. '
    'Tính năng Dashboard điều hành cung cấp cái nhìn tổng thể, hỗ trợ ban lãnh đạo ra quyết định nhanh chóng '
    'và chính xác hơn.')

add_body_paragraph(doc,
    'Tuy nhiên, khóa luận cũng nhận thấy một số hạn chế cần được khắc phục trong tương lai như: giao diện '
    'chưa tối ưu hoàn toàn cho thiết bị di động, chưa áp dụng công nghệ mã QR/Barcode trong kiểm kê thực tế, '
    'và tính năng AI mới dừng ở mức hỗ trợ cơ bản. Những hạn chế này đã được đề xuất giải pháp khắc phục '
    'tại mục hướng phát triển.')

add_body_paragraph(doc,
    'Có thể khẳng định rằng, việc ứng dụng phần mềm GlassFlow đã giúp Công ty CP ViralWindow chuyển đổi từ '
    'phương thức quản lý kho truyền thống sang quản lý số hóa toàn diện, góp phần tối ưu hóa nguồn lực, '
    'giảm thiểu thất thoát và nâng cao hiệu quả điều hành doanh nghiệp. Kết quả nghiên cứu không chỉ có '
    'giá trị thực tiễn đối với ViralWindow mà còn có thể tham khảo áp dụng cho các doanh nghiệp cùng ngành '
    'sản xuất và kinh doanh cửa nhôm kính.')

# ===================== SAVE =====================
doc.save(OUTPUT_PATH)
print(f"\n✅ Đã lưu file thành công tại:\n{OUTPUT_PATH}")
print("Các mục đã thêm: 4.4 (5 tiểu mục + bảng), 4.5 (6 hướng), KẾT LUẬN")
