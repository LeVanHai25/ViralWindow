import docx
import sys
sys.stdout.reconfigure(encoding='utf-8')

filepath = r'd:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main.docx'
doc = docx.Document(filepath)

# Check page setup details
print('=== PAGE SETUP ===')
for i, section in enumerate(doc.sections):
    print(f'Section {i+1}:')
    print(f'  Page: {round(section.page_width.cm,2)} x {round(section.page_height.cm,2)} cm')
    print(f'  Left: {round(section.left_margin.cm,2)} cm | Right: {round(section.right_margin.cm,2)} cm')
    print(f'  Top: {round(section.top_margin.cm,2)} cm | Bottom: {round(section.bottom_margin.cm,2)} cm')

# Check fonts
print('\n=== TOP FONTS ===')
font_stats = {}
for para in doc.paragraphs:
    if para.text.strip():
        for run in para.runs:
            fn = run.font.name or 'Inherited'
            fs = f'{run.font.size.pt}pt' if run.font.size else 'Inherited'
            bold = 'B' if run.font.bold else ''
            italic = 'I' if run.font.italic else ''
            key = f'{fn} | {fs} | {bold}{italic}'
            font_stats[key] = font_stats.get(key, 0) + 1

for k, v in sorted(font_stats.items(), key=lambda x: -x[1])[:25]:
    print(f'  {k}: {v} runs')

# Check Heading numbering
print('\n=== HEADING STRUCTURE ===')
for para in doc.paragraphs:
    sn = para.style.name
    t = para.text.strip()
    if sn == 'Heading 1' and t:
        print(f'[H1] "{t}"')
    elif sn == 'Heading 2' and t:
        print(f'  [H2] "{t}"')
    elif sn == 'Heading 3' and t:
        print(f'    [H3] "{t}"')

# Check header/footer
print('\n=== HEADERS/FOOTERS ===')
for i, section in enumerate(doc.sections):
    h = section.header
    f = section.footer
    if h and h.paragraphs:
        ht = ' | '.join([p.text for p in h.paragraphs if p.text.strip()])
        if ht:
            print(f'  Section {i+1} Header: {ht[:80]}')
    if f and f.paragraphs:
        ft = ' | '.join([p.text for p in f.paragraphs if p.text.strip()])
        if ft:
            print(f'  Section {i+1} Footer: {ft[:80]}')

# Content before Chapter 1 - check for required sections
print('\n=== PRE-CHAPTER CONTENT ===')
for para in doc.paragraphs:
    t = para.text.strip()
    if para.style.name == 'Heading 1' and 'CHƯƠNG 1' in t.upper():
        break
    if t:
        print(f'  [{para.style.name}] {t[:100]}')

# Check TOC
print('\n=== TOC STYLES ===')
toc_count = 0
for para in doc.paragraphs:
    if 'toc' in para.style.name.lower():
        toc_count += 1
        if toc_count <= 10:
            print(f'  [{para.style.name}] {para.text.strip()[:80]}')
print(f'  Total TOC entries: {toc_count}')

# References section
print('\n=== REFERENCES ===')
in_ref = False
for para in doc.paragraphs:
    t = para.text.strip()
    if 'TÀI LIỆU THAM KHẢO' in t.upper():
        in_ref = True
        print(f'  [{para.style.name}] {t}')
        continue
    if in_ref and t:
        print(f'  [{para.style.name}] {t[:120]}')
