import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx():
    # Setup document
    doc = Document()
    
    # HaUI 2025 Standard Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)
        # Header/Footer distance
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.27)

    # Style Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(6)

    drafts_dir = r"d:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\Drafts"
    files = [
        "01_Mo_Dau_Chuong_1.md",
        "02_Chuong_2_Phan_Tich.md",
        "03_Chuong_3_Thiet_Ke.md",
        "04_Chuong_4_Cai_Dat_Kiem_Thu.md",
        "05_Ket_Luan_Tai_Lieu.md"
    ]

    for filename in files:
        filepath = os.path.join(drafts_dir, filename)
        if not os.path.exists(filepath):
            continue
            
        with open(filepath, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Heading Level 1 (# )
            if line.startswith('# '):
                h = doc.add_heading(line[2:], level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set font for heading manually as add_heading resets it
                for run in h.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run.font.color.rgb = None # Black
                    
            # Heading Level 2 (## ) or Level 3 (### )
            elif line.startswith('### '):
                h = doc.add_heading(line[4:], level=3)
                for run in h.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = None
            elif line.startswith('## '):
                h = doc.add_heading(line[3:], level=2)
                for run in h.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = None
                    
            # Paragraphs
            else:
                # Clean markdown bold/italic
                line = line.replace('**', '').replace('__', '')
                p = doc.add_paragraph(line)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Page break after each chapter except the last
        if filename != files[-1]:
            doc.add_page_break()

    output_path = r"d:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\Do_An_Tot_Nghiep_LeVanHai.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    generate_docx()
