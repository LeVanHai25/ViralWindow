# -*- coding: utf-8 -*-
"""Kiem tra toan bo noi dung FINAL.docx - bao cao cuoi cung"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

for root,dirs,files in os.walk(r'D:\ViralWindow_Phan_Mem_Nhom_Kinh'):
    for f in files:
        if 'FINAL' in f and f.endswith('.docx') and not f.startswith('~'):
            path = os.path.join(root, f)

doc = Document(path)
all_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# 1. Kiem tra cac muc da them
MUST_HAVE = {
    'DANH MUC CAC TU VIET TAT': 'DANH MỤC CÁC TỪ VIẾT TẮT',
    'BOM Engine':               '3.3.4. Thuật toán Bóc tách Vật tư (BOM Engine)',
    'requirePermission':        'Code RBAC middleware',
    '4.2.9':                    'Kiểm thử bảo mật',
    'optimizeCutting':          'Code thuật toán FFD',
    'So sanh voi':              'Bảng so sánh hệ thống',
}
found = {}
for key, label in MUST_HAVE.items():
    found[label] = any(key.lower() in t.lower() for t in all_text)
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for key, label in MUST_HAVE.items():
                if key.lower() in cell.text.lower():
                    found[label] = True

# 2. Kiem tra danh so muc Chuong 3
c3_nums = [t for t in all_text if len(t)<75 and t and t[0]=='3' and '.' in t[:5]]

# 3. Kiem tra le trang
section = doc.sections[0]
margin_top = round(section.top_margin.cm, 1) if section.top_margin else 0
margin_bot = round(section.bottom_margin.cm, 1) if section.bottom_margin else 0
margin_left = round(section.left_margin.cm, 1) if section.left_margin else 0
margin_right = round(section.right_margin.cm, 1) if section.right_margin else 0

# 4. Kiem tra tu viet tat
abbr_count = sum(1 for tbl in doc.tables[:5]
                 for row in tbl.rows
                 for cell in row.cells[:1]
                 if len(cell.text.strip()) <= 10 and cell.text.strip().isupper())

print("=" * 60)
print("BAO CAO KIEM TRA CUOI CUNG - FINAL.docx")
print("=" * 60)
print(f"\nTong so doan van : {len(doc.paragraphs)}")
print(f"Tong so bang     : {len(doc.tables)}")
print(f"\nLE TRANG:")
print(f"  Tren  : {margin_top}cm  {'OK' if margin_top >= 2.9 else 'SAI'}")
print(f"  Duoi  : {margin_bot}cm  {'OK' if margin_bot >= 2.9 else 'SAI'}")
print(f"  Trai  : {margin_left}cm  {'OK' if margin_left >= 3.4 else 'SAI'}")
print(f"  Phai  : {margin_right}cm  {'OK' if margin_right >= 1.9 else 'SAI'}")
print(f"\nNOI DUNG THEM MOI:")
for label, ok in found.items():
    print(f"  [{'OK' if ok else 'MISSING'}] {label}")
print(f"\nDANH SO MUC CHUONG 3:")
for t in c3_nums:
    print(f"  {t}")
print(f"\nSO TU VIET TAT (uoc tinh): {abbr_count}")
print("\n" + "=" * 60)
ok_count = sum(1 for v in found.values() if v)
total = len(found)
le_ok = all([margin_top>=2.9, margin_bot>=2.9, margin_left>=3.4, margin_right>=1.9])
print(f"KET QUA: {ok_count}/{total} noi dung moi co mat")
print(f"LE TRANG: {'DAT CHUAN' if le_ok else 'CAN KIEM TRA'}")
if ok_count == total and le_ok:
    print("\n✅ FILE FINAL DAT CHUAN - SAN SANG NOP!")
else:
    print("\n⚠️  CON MOT SO MUC CAN KIEM TRA THU CONG TRONG WORD")
