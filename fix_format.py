# -*- coding: utf-8 -*-
"""Tạo báo cáo HTML kiểm tra định dạng đồ án và sửa các lỗi có thể sửa tự động"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_v2.docx"
out_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"

doc = Document(doc_path)
fixes = []

# ==========================================================
# FIX 1: Sửa lề trang - trên 3cm, dưới 3cm
# ==========================================================
for section in doc.sections:
    old_top = section.top_margin.cm if section.top_margin else 0
    old_bot = section.bottom_margin.cm if section.bottom_margin else 0
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    fixes.append(f"✅ Lề trên: {old_top:.1f}cm → 3.0cm")
    fixes.append(f"✅ Lề dưới: {old_bot:.1f}cm → 3.0cm")

# ==========================================================
# FIX 2: Sửa thụt đầu dòng sai (0.81cm → 1.25cm)
# ==========================================================
HEADING_STYLES = {"heading 1","heading 2","heading 3","heading 4","heading 5"}
HEADING_KEYS   = ["CHƯƠNG","MỞ ĐẦU","KẾT LUẬN","TÀI LIỆU THAM KHẢO","DANH MỤC","MỤC LỤC"]

indent_fixed = 0
for para in doc.paragraphs:
    txt = para.text.strip()
    if not txt or len(txt) < 10:
        continue
    style_name = para.style.name.lower()
    is_heading = any(s in style_name for s in HEADING_STYLES) or \
                 any(k in txt.upper() for k in HEADING_KEYS)
    if is_heading:
        continue

    pf = para.paragraph_format
    fi = pf.first_line_indent
    if fi is not None:
        fi_cm = fi / 360000
        if 0 < fi_cm < 1.0:  # indent sai (< 1cm)
            pf.first_line_indent = Cm(1.25)
            indent_fixed += 1

fixes.append(f"✅ Sửa {indent_fixed} đoạn có thụt đầu dòng sai → 1.25cm")

# ==========================================================
# FIX 3: Sửa các đoạn chưa căn đều (Justify) - chỉ thân văn
# ==========================================================
align_fixed = 0
SKIP_ALIGNS = ["Hình ", "Bảng ", "Nguồn:", "Source:"]

for para in doc.paragraphs:
    txt = para.text.strip()
    if not txt or len(txt) < 20:
        continue
    style_name = para.style.name.lower()
    is_heading = any(s in style_name for s in HEADING_STYLES) or \
                 any(k in txt.upper() for k in HEADING_KEYS)
    if is_heading:
        continue
    if any(txt.startswith(sk) for sk in SKIP_ALIGNS):
        continue

    # Chỉ sửa nếu đang CENTER hoặc RIGHT
    align = para.paragraph_format.alignment
    if align in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        align_fixed += 1

fixes.append(f"✅ Sửa {align_fixed} đoạn chưa căn đều hai bên → Justify")

# ==========================================================
# FIX 4: Đảm bảo heading cấp 1 (CHƯƠNG X) in đậm
# ==========================================================
heading_fixed = 0
CHAPTER_KEYWORDS = ["CHƯƠNG 1","CHƯƠNG 2","CHƯƠNG 3","CHƯƠNG 4",
                    "MỞ ĐẦU","KẾT LUẬN","TÀI LIỆU THAM KHẢO",
                    "DANH MỤC BẢNG","DANH MỤC HÌNH"]

for para in doc.paragraphs:
    txt = para.text.strip().upper()
    if any(kw in txt for kw in CHAPTER_KEYWORDS) and len(para.text.strip()) < 80:
        for run in para.runs:
            if not run.bold:
                run.bold = True
                heading_fixed += 1
                break

fixes.append(f"✅ Đặt in đậm cho {heading_fixed} tiêu đề chương chính")

# ==========================================================
# Lưu file FINAL
# ==========================================================
doc.save(out_path)
print("FIXES APPLIED:")
for f in fixes:
    print(" ", f)
print(f"\n✅ Đã lưu: {out_path}")
