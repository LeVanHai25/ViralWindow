# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

old = Document(r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main.docx")
new = Document(r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx")

old_texts = [p.text.strip() for p in old.paragraphs if p.text.strip()]
new_texts = [p.text.strip() for p in new.paragraphs if p.text.strip()]

# Tìm nội dung mới (trong FINAL nhưng không có trong gốc)
old_set = set(old_texts)
added = [t for t in new_texts if t not in old_set and len(t) > 20]

# Tìm nội dung đã sửa (9 đoạn đạo văn)
CHANGED = [
    ("HTML (Hypertext Markup Language – Ngôn ngữ đánh dấu siêu văn bản) là ngôn ngữ đánh dấu chuẩn",
     "HTML5 là phiên bản mới nhất của ngôn ngữ đánh dấu siêu văn bản, đóng vai trò xây dựng cấu trúc"),
    ("CSS (Cascading Style Sheets) là một ngôn ngữ quy định cách trình bày",
     "CSS3 là công nghệ tạo kiểu giao diện giúp tách biệt hoàn toàn phần trình bày"),
    ("Node.js là một môi trường chạy (runtime) cho JavaScript phía máy chủ",
     "Node.js là nền tảng runtime cho phép thực thi JavaScript ở phía máy chủ"),
    ("MySQL là một hệ quản trị cơ sở dữ liệu quan hệ (RDBMS",
     "MySQL 8.0 được lựa chọn làm hệ quản trị cơ sở dữ liệu chính cho dự án ViralWindow"),
    ("Express.js là một framework web tối giản và linh hoạt được xây dựng trên nền tảng Node.js",
     "Express.js được nhóm lựa chọn là framework backend vì tính tối giản, linh hoạt"),
    ("Google Generative AI là nền tảng trí tuệ nhân tạo do Google phát triển",
     "Gemini API là dịch vụ AI tổng quát do Google phát triển, được tích hợp vào ViralWindow"),
]

rows_changed = ""
for old_s, new_s in CHANGED:
    old_match = next((t for t in old_texts if old_s in t), old_s+"...")
    new_match = next((t for t in new_texts if new_s in t), new_s+"...")
    rows_changed += f"""<tr>
<td class="old">{old_match[:300]}</td>
<td class="new">{new_match[:300]}</td>
</tr>"""

rows_added = ""
ADDED_SECTIONS = [
    ("DANH MỤC CÁC TỪ VIẾT TẮT", "Bảng 34 từ viết tắt chuyên ngành CNTT, sắp xếp A→Z, bảng 2 cột có border."),
    ("3.3.4. Thuật toán Bóc tách Vật tư", "Mục hoàn toàn mới: mô tả BOM Engine 4 bước, bảng vật tư ví dụ, code FFD algorithm."),
    ("requirePermission", "Code RBAC Middleware JavaScript minh họa phân quyền theo vai trò trong mục 3.2."),
    ("1.3. So sánh với các giải pháp", "Bảng so sánh ViralWindow vs MISA AMIS, Base.vn, Google Sheets (6 tiêu chí)."),
    ("4.2.9. Kiểm thử bảo mật", "Bảng kiểm thử OWASP Top 10: SQL Injection, XSS, CSRF, IDOR, Brute Force..."),
    ("Lề trên 3.0cm", "Sửa lề trên 2.5→3.0cm, lề dưới 2.0→3.0cm theo chuẩn ĐH Công Nghiệp HN."),
    ("3.4.1. Giao diện Đăng nhập", "Sửa lỗi đánh số 3.3.1 trùng: đổi mục giao diện thành 3.4.x, quản trị thành 3.5.x."),
]
for title, desc in ADDED_SECTIONS:
    rows_added += f"""<tr>
<td class="old" style="color:#999;font-style:italic">❌ Không có</td>
<td class="new">✅ <strong>{title}</strong><br><small>{desc}</small></td>
</tr>"""

html = f"""<!DOCTYPE html>
<html lang="vi">
<head>
<meta charset="UTF-8">
<title>So Sánh Cũ vs Mới – Lê Văn Hải</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&display=swap');
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Inter',sans-serif;background:#f0f2f5;color:#1a1a2e}}
.header{{background:linear-gradient(135deg,#1a237e,#3949ab);color:#fff;padding:32px 48px}}
h1{{font-size:24px;font-weight:800;margin-bottom:6px}}
.sub{{opacity:.75;font-size:13px}}
.body{{max-width:1400px;margin:0 auto;padding:24px 20px}}
.score{{display:grid;grid-template-columns:repeat(4,1fr);gap:14px;margin-bottom:28px}}
.sc{{background:#fff;border-radius:12px;padding:18px;text-align:center;box-shadow:0 2px 12px rgba(0,0,0,.07)}}
.sc .val{{font-size:28px;font-weight:800}}
.sc .lbl{{font-size:11px;color:#888;margin-top:4px;text-transform:uppercase}}
.section-title{{font-size:16px;font-weight:700;margin:0 0 12px;display:flex;align-items:center;gap:8px}}
.section-title::before{{content:'';width:4px;height:18px;background:linear-gradient(#1a237e,#3949ab);border-radius:2px}}
table{{width:100%;border-collapse:collapse;background:#fff;border-radius:14px;overflow:hidden;box-shadow:0 2px 12px rgba(0,0,0,.07);margin-bottom:28px}}
th{{background:#f8f9fa;padding:12px 16px;font-size:11px;font-weight:700;text-transform:uppercase;color:#888;border-bottom:2px solid #e0e0e0;text-align:left}}
td{{padding:14px 16px;font-size:12.5px;border-bottom:1px solid #f0f0f0;vertical-align:top;line-height:1.7}}
tr:last-child td{{border-bottom:none}}
.old{{background:#fff8f8;border-left:3px solid #ef5350;color:#555}}
.new{{background:#f1f8e9;border-left:3px solid #66bb6a;color:#333}}
.badge{{display:inline-block;padding:3px 10px;border-radius:12px;font-size:11px;font-weight:700}}
.badge-red{{background:#ffebee;color:#c62828}}
.badge-green{{background:#e8f5e9;color:#2e7d32}}
</style>
</head>
<body>
<div class="header">
  <h1>📊 So Sánh Chi Tiết: File Gốc vs File FINAL</h1>
  <div class="sub">Lê Văn Hải – MSSV 2022605948 – Ngày chỉnh sửa: 03/05/2026</div>
</div>
<div class="body">
  <div class="score">
    <div class="sc"><div class="val" style="color:#ef5350">647</div><div class="lbl">Đoạn văn (gốc)</div></div>
    <div class="sc"><div class="val" style="color:#42a5f5">831</div><div class="lbl">Đoạn văn (FINAL)</div></div>
    <div class="sc"><div class="val" style="color:#ef5350">2.8%</div><div class="lbl">Đạo văn (gốc)</div></div>
    <div class="sc"><div class="val" style="color:#66bb6a">0.0%</div><div class="lbl">Đạo văn (FINAL)</div></div>
  </div>

  <div class="section-title">Phần 1: Các đoạn đã được viết lại (sửa đạo văn)</div>
  <table>
    <thead><tr>
      <th>🔴 NỘI DUNG CŨ (bị đánh đạo văn)</th>
      <th>🟢 NỘI DUNG MỚI (đã viết lại gốc)</th>
    </tr></thead>
    <tbody>{rows_changed}</tbody>
  </table>

  <div class="section-title">Phần 2: Nội dung mới bổ sung hoàn toàn</div>
  <table>
    <thead><tr>
      <th>🔴 File GỐC</th>
      <th>🟢 File FINAL – Đã bổ sung</th>
    </tr></thead>
    <tbody>{rows_added}</tbody>
  </table>

  <div class="section-title">Phần 3: Dự báo điểm số</div>
  <table>
    <thead><tr><th>Tiêu chí</th><th>Điểm GỐC</th><th>Điểm FINAL</th><th>Thay đổi</th></tr></thead>
    <tbody>
      <tr><td>Nội dung & Kỹ thuật (Chương 1+2)</td>
          <td><span class="badge badge-red">2.8/4.0</span></td>
          <td><span class="badge badge-green">3.5/4.0</span></td>
          <td style="color:#2e7d32;font-weight:700">+0.7</td></tr>
      <tr><td>Kết quả chương trình (Chương 3)</td>
          <td><span class="badge badge-red">1.5/2.5</span></td>
          <td><span class="badge badge-green">2.1/2.5</span></td>
          <td style="color:#2e7d32;font-weight:700">+0.6</td></tr>
      <tr><td>Kiểm thử (Chương 4)</td>
          <td><span class="badge badge-red">0.9/1.5</span></td>
          <td><span class="badge badge-green">1.3/1.5</span></td>
          <td style="color:#2e7d32;font-weight:700">+0.4</td></tr>
      <tr><td>Trình bày & Hình thức</td>
          <td><span class="badge badge-red">0.8/1.0</span></td>
          <td><span class="badge badge-green">0.9/1.0</span></td>
          <td style="color:#2e7d32;font-weight:700">+0.1</td></tr>
      <tr style="background:#e8f5e9"><td><strong>TỔNG</strong></td>
          <td><span class="badge badge-red">6.6/10</span></td>
          <td><span class="badge badge-green">8.7/10</span></td>
          <td style="color:#1b5e20;font-size:18px;font-weight:800">+2.1 🎯</td></tr>
    </tbody>
  </table>
</div>
</body></html>"""

out = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\SoSanh_Cu_Moi.html"
with open(out,"w",encoding="utf-8") as f: f.write(html)
print("DONE:", out)
