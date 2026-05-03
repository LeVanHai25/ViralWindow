# -*- coding: utf-8 -*-
"""
Tạo DANH MỤC CÁC TỪ VIẾT TẮT dạng BẢNG 2 CỘT
Đúng mẫu chuẩn ĐH Công Nghiệp Hà Nội:
 - Tiêu đề: DANH MỤC CÁC TỪ VIẾT TẮT (in đậm, căn giữa)
 - Bảng 2 cột: Cột 1 = Từ viết tắt (căn giữa), Cột 2 = Giải thích
"""
import sys, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

src_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_v2.docx"
out_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"

# Danh sách từ viết tắt – cột 2 chỉ ghi giải thích ngắn gọn (như mẫu)
ABBREVIATIONS = sorted([
    ("2FA",    "Xác thực hai yếu tố (Two-Factor Authentication)"),
    ("API",    "Giao diện lập trình ứng dụng (Application Programming Interface)"),
    ("BOM",    "Danh mục vật tư (Bill of Materials)"),
    ("CORS",   "Chia sẻ tài nguyên giữa các nguồn (Cross-Origin Resource Sharing)"),
    ("CRUD",   "Tạo, Đọc, Cập nhật, Xóa (Create, Read, Update, Delete)"),
    ("CSRF",   "Giả mạo yêu cầu liên trang (Cross-Site Request Forgery)"),
    ("CSS",    "Ngôn ngữ định kiểu tầng (Cascading Style Sheets)"),
    ("CSDL",   "Cơ sở dữ liệu"),
    ("DCL",    "Ngôn ngữ kiểm soát dữ liệu (Data Control Language)"),
    ("DDL",    "Ngôn ngữ định nghĩa dữ liệu (Data Definition Language)"),
    ("ERP",    "Hoạch định nguồn lực doanh nghiệp (Enterprise Resource Planning)"),
    ("HTML",   "Ngôn ngữ đánh dấu siêu văn bản (HyperText Markup Language)"),
    ("HTTP",   "Giao thức truyền tải siêu văn bản (HyperText Transfer Protocol)"),
    ("HTTPS",  "Giao thức HTTP bảo mật (HyperText Transfer Protocol Secure)"),
    ("JWT",    "Mã thông báo xác thực JSON (JSON Web Token)"),
    ("KPI",    "Chỉ số hiệu suất chính (Key Performance Indicator)"),
    ("MySQL",  "Hệ quản trị cơ sở dữ liệu quan hệ mã nguồn mở"),
    ("npm",    "Trình quản lý gói Node.js (Node Package Manager)"),
    ("OWASP",  "Dự án bảo mật ứng dụng web mở (Open Web Application Security Project)"),
    ("PDF",    "Định dạng tài liệu di động (Portable Document Format)"),
    ("RBAC",   "Kiểm soát truy cập theo vai trò (Role-Based Access Control)"),
    ("RDBMS",  "Hệ quản trị cơ sở dữ liệu quan hệ (Relational Database Management System)"),
    ("REST",   "Kiến trúc truyền trạng thái đại diện (Representational State Transfer)"),
    ("RWD",    "Thiết kế web thích ứng (Responsive Web Design)"),
    ("SEO",    "Tối ưu hóa công cụ tìm kiếm (Search Engine Optimization)"),
    ("SQL",    "Ngôn ngữ truy vấn có cấu trúc (Structured Query Language)"),
    ("SSL",    "Lớp kết nối bảo mật (Secure Sockets Layer)"),
    ("TCL",    "Ngôn ngữ kiểm soát giao dịch (Transaction Control Language)"),
    ("TLS",    "Bảo mật tầng vận chuyển (Transport Layer Security)"),
    ("UML",    "Ngôn ngữ mô hình hóa thống nhất (Unified Modeling Language)"),
    ("URL",    "Định vị tài nguyên thống nhất (Uniform Resource Locator)"),
    ("UX",     "Trải nghiệm người dùng (User Experience)"),
    ("XHTML",  "Ngôn ngữ HTML mở rộng (Extensible HyperText Markup Language)"),
    ("XSS",    "Tấn công kịch bản liên trang (Cross-Site Scripting)"),
], key=lambda x: x[0].upper())

# ─────────────────────────────────────────────
# Hàm thiết lập border cho bảng
# ─────────────────────────────────────────────
def set_table_borders(table):
    """Đặt border đơn cho toàn bộ bảng (như mẫu ảnh)"""
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_margins(cell, top=80, bottom=80, left=108, right=108):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top),('bottom', bottom),('left', left),('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def fmt_run(run, size=13, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold

# ─────────────────────────────────────────────
# Load tài liệu gốc (v2)
# ─────────────────────────────────────────────
doc = Document(src_path)

# Sửa lề
for section in doc.sections:
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(3.0)

# Tìm vị trí MỞ ĐẦU
mo_dau_idx = None
for i, para in enumerate(doc.paragraphs):
    if "MỞ ĐẦU" in para.text.strip().upper() and len(para.text.strip()) < 20:
        mo_dau_idx = i
        break

print(f"MỞ ĐẦU tại index: {mo_dau_idx}")
mo_dau_elem = doc.paragraphs[mo_dau_idx]._element

# ─────────────────────────────────────────────
# Tạo các element chèn vào (theo thứ tự đúng)
# Dùng Document tạm để build rồi addprevious
# ─────────────────────────────────────────────

def insert_elem_before(ref_elem, elem):
    ref_elem.addprevious(elem)

# 1. Page break trước mục
pb = OxmlElement('w:p')
pbPr = OxmlElement('w:pPr')
pb_r = OxmlElement('w:r')
pb_br = OxmlElement('w:br')
pb_br.set(qn('w:type'), 'page')
pb_r.append(pb_br)
pb.append(pbPr)
pb.append(pb_r)
insert_elem_before(mo_dau_elem, pb)

# 2. Tiêu đề "DANH MỤC CÁC TỪ VIẾT TẮT"
tmp = Document()
p_title = tmp.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(12)
r_title = p_title.add_run("DANH MỤC CÁC TỪ VIẾT TẮT")
fmt_run(r_title, size=13, bold=True)
insert_elem_before(mo_dau_elem, copy.deepcopy(p_title._element))

# 3. Bảng 2 cột
# Tạo bảng trong tmp doc rồi move element vào
tmp2 = Document()
table = tmp2.add_table(rows=len(ABBREVIATIONS), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Độ rộng cột: cột 1 = 3.5cm, cột 2 = phần còn lại (~12.5cm)
from docx.shared import Cm
col_widths = [Cm(3.5), Cm(12.5)]
for i, row in enumerate(table.rows):
    abbr, meaning = ABBREVIATIONS[i]

    # Cột 1: Từ viết tắt – căn giữa, in đậm
    cell1 = row.cells[0]
    cell1.width = col_widths[0]
    set_cell_margins(cell1, top=60, bottom=60, left=80, right=80)
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(2)
    r1 = p1.add_run(abbr)
    fmt_run(r1, size=13, bold=False)

    # Cột 2: Giải thích – căn trái
    cell2 = row.cells[1]
    cell2.width = col_widths[1]
    set_cell_margins(cell2, top=60, bottom=60, left=120, right=80)
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(2)
    r2 = p2.add_run(meaning)
    fmt_run(r2, size=13, bold=False)

# Đặt border bảng
set_table_borders(table)

# Set chiều rộng cột
for row in table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(12.5)

# Copy tbl element vào document chính
tbl_elem = copy.deepcopy(table._tbl)
insert_elem_before(mo_dau_elem, tbl_elem)

# 4. Dòng trống sau bảng
tmp3 = Document()
p_blank = tmp3.add_paragraph()
p_blank.paragraph_format.space_before = Pt(0)
p_blank.paragraph_format.space_after  = Pt(0)
insert_elem_before(mo_dau_elem, copy.deepcopy(p_blank._element))

# ─────────────────────────────────────────────
# Lưu
# ─────────────────────────────────────────────
doc.save(out_path)
print(f"✅ Đã tạo bảng 2 cột với {len(ABBREVIATIONS)} từ viết tắt")
print(f"✅ Lưu: {out_path}")

# Xác minh
doc2 = Document(out_path)
tables = doc2.tables
print(f"✅ Số bảng trong file: {len(tables)}")
if tables:
    t = tables[0]
    print(f"✅ Bảng đầu tiên: {len(t.rows)} hàng x {len(t.columns)} cột")
    for row in t.rows[:5]:
        c1 = row.cells[0].text.strip()
        c2 = row.cells[1].text.strip()[:50]
        print(f"   | {c1:<8} | {c2}")
