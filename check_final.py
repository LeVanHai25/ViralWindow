# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"
doc = Document(path)

checks = ['3.3.4', 'BOM', 'DANH MUC', 'requirePermission', '4.2.9', 'optimizeCutting', 'CHUC_VU_QUYEN']
found = {c: False for c in checks}

all_text = []
for p in doc.paragraphs:
    t = p.text.strip()
    if t:
        all_text.append(t)
    for c in checks:
        if c.lower() in t.lower():
            found[c] = True

for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            t = cell.text.strip()
            for c in checks:
                if c.lower() in t.lower():
                    found[c] = True

print("=== KIEM TRA NOI DUNG FINAL ===")
print(f"Tong doan van: {len(doc.paragraphs)}")
print(f"Tong bang: {len(doc.tables)}")
print()
print("=== TIM KIEM NOI DUNG ===")
for k, v in found.items():
    status = "CO" if v else "KHONG CO"
    print(f"  [{status}] '{k}'")

# Xuat 50 dong dau tien de kiem tra
print()
print("=== 30 DONG DAU TIEN ===")
for i, t in enumerate(all_text[:30]):
    print(f"[{i}] {t[:100]}")

# Tim cac muc 3.x de kiem tra danh so
print()
print("=== DANH SO MUC CHUONG 3 ===")
for t in all_text:
    if t.startswith('3.') and len(t) < 80:
        print(f"  {t}")
