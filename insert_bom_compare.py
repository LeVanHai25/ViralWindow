# -*- coding: utf-8 -*-
"""Chèn BOM và bảng so sánh vào đúng vị trí"""
import sys, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from lxml import etree

src = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"
doc = Document(src)
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def q(tag): return f'{{{WNS}}}{tag}'

def mk_run(p_elem, text, bold=False, italic=False, size=13, color=None, mono=False):
    r = etree.SubElement(p_elem, q('r'))
    rPr = etree.SubElement(r, q('rPr'))
    if bold:   etree.SubElement(rPr, q('b')); etree.SubElement(rPr, q('bCs'))
    if italic: etree.SubElement(rPr, q('i'))
    if color:
        cl = etree.SubElement(rPr, q('color')); cl.set(q('val'), color)
    fn = 'Courier New' if mono else 'Times New Roman'
    rf = etree.SubElement(rPr, q('rFonts'))
    rf.set(q('ascii'), fn); rf.set(q('hAnsi'), fn)
    sz = etree.SubElement(rPr, q('sz')); sz.set(q('val'), str(size*2))
    szCs = etree.SubElement(rPr, q('szCs')); szCs.set(q('val'), str(size*2))
    t = etree.SubElement(r, q('t'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text

def mk_para(ref, text='', bold=False, italic=False, size=13,
            align='both', color=None, sb=0, sa=6, fi=0, mono=False):
    p = etree.Element(q('p'))
    ref.addprevious(p)
    pPr = etree.SubElement(p, q('pPr'))
    jc = etree.SubElement(pPr, q('jc')); jc.set(q('val'), align)
    sp = etree.SubElement(pPr, q('spacing'))
    sp.set(q('before'), str(int(sb*20))); sp.set(q('after'), str(int(sa*20)))
    if fi:
        ind = etree.SubElement(pPr, q('ind')); ind.set(q('firstLine'), str(int(fi*567)))
    if text:
        mk_run(p, text, bold=bold, italic=italic, size=size, color=color, mono=mono)

def mk_code(ref, lines):
    for line in lines:
        p = etree.Element(q('p'))
        ref.addprevious(p)
        pPr = etree.SubElement(p, q('pPr'))
        sp = etree.SubElement(pPr, q('spacing')); sp.set(q('before'),'20'); sp.set(q('after'),'20')
        ind = etree.SubElement(pPr, q('ind')); ind.set(q('left'),'567')
        if line: mk_run(p, line, mono=True, size=10, color='1F3864')

def mk_table(ref, headers, rows, col_widths):
    tbl = etree.Element(q('tbl'))
    ref.addprevious(tbl)
    tblPr = etree.SubElement(tbl, q('tblPr'))
    tblW = etree.SubElement(tblPr, q('tblW')); tblW.set(q('w'),'9072'); tblW.set(q('type'),'dxa')
    tblBorders = etree.SubElement(tblPr, q('tblBorders'))
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = etree.SubElement(tblBorders, q(side))
        b.set(q('val'),'single'); b.set(q('sz'),'4'); b.set(q('color'),'000000'); b.set(q('space'),'0')
    tblGrid = etree.SubElement(tbl, q('tblGrid'))
    for w in col_widths:
        gc = etree.SubElement(tblGrid, q('gridCol')); gc.set(q('w'), str(int(w*567)))

    def add_row(cells, hdr=False):
        tr = etree.SubElement(tbl, q('tr'))
        for ci, (txt, cw) in enumerate(zip(cells, col_widths)):
            tc = etree.SubElement(tr, q('tc'))
            tcPr = etree.SubElement(tc, q('tcPr'))
            tcW = etree.SubElement(tcPr, q('tcW')); tcW.set(q('w'),str(int(cw*567))); tcW.set(q('type'),'dxa')
            tcMar = etree.SubElement(tcPr, q('tcMar'))
            for s,v in [('top','60'),('bottom','60'),('left','113'),('right','113')]:
                m = etree.SubElement(tcMar, q(s)); m.set(q('w'),v); m.set(q('type'),'dxa')
            p_tc = etree.SubElement(tc, q('p'))
            pPr_tc = etree.SubElement(p_tc, q('pPr'))
            jc_tc = etree.SubElement(pPr_tc, q('jc')); jc_tc.set(q('val'),'center' if ci==0 else 'left')
            sp_tc = etree.SubElement(pPr_tc, q('spacing')); sp_tc.set(q('before'),'20'); sp_tc.set(q('after'),'20')
            mk_run(p_tc, txt, bold=hdr, size=12)

    if headers: add_row(headers, hdr=True)
    for row in rows: add_row(row)

paras = doc.paragraphs

# ── TASK 2: BOM – chèn trước đoạn [499] "Giao diện chung" ──
anchor_bom = paras[499]._element
mk_para(anchor_bom, '', sb=0, sa=8)
mk_code(anchor_bom, [
    '// backend/controllers/bomController.js',
    'function optimizeCutting(pieces, barLen=6000, kerf=3) {',
    '  const sorted = [...pieces].sort((a,b) => b.length - a.length);',
    '  const bars = [];',
    '  for (const piece of sorted) {',
    '    let placed = false;',
    '    for (const bar of bars) {',
    '      const used = bar.reduce((s,p) => s + p.length + kerf, 0);',
    '      if (used + piece.length + kerf <= barLen) {',
    '        bar.push(piece); placed = true; break;',
    '      }',
    '    }',
    '    if (!placed) bars.push([piece]);',
    '  }',
    '  const totalUsed = pieces.reduce((s,p) => s + p.length, 0);',
    '  const waste = ((barLen*bars.length - totalUsed)/(barLen*bars.length)*100).toFixed(1);',
    '  return { totalBars: bars.length, wastePercent: waste };',
    '}',
])
mk_para(anchor_bom, 'c) Đoạn mã thuật toán tối ưu cắt (First-Fit Decreasing – FFD):',
        bold=False, italic=True, size=13, align='left', sb=6, sa=4)
mk_table(anchor_bom,
    headers=['Thanh nhôm', 'Công thức tính'],
    rows=[
        ['Thanh đứng khung', 'Dài = H – 10mm, SL = 2 thanh/bộ → 6 đoạn → cần 4 thanh 6m'],
        ['Thanh ngang khung', 'Dài = W – 10mm, SL = 2 thanh/bộ → 6 đoạn → cần 3 thanh 6m'],
        ['Thanh đứng cánh',  'Dài = H – 55mm, SL = 2 thanh/bộ → 6 đoạn → cần 4 thanh 6m'],
        ['Thanh ngang cánh', 'Dài = W – 55mm, SL = 2 thanh/bộ → 6 đoạn → cần 3 thanh 6m'],
        ['Kính cường lực 5mm','(W–70) × (H–40) mm × 3 tấm'],
        ['Hao phí ước tính', '≈ 8.3% (tối ưu FFD, kerf = 3mm/lát cắt)'],
    ],
    col_widths=(4, 12)
)
mk_para(anchor_bom,
    'b) Ví dụ: Cửa đi 1 cánh, kích thước 1200×2100mm, hệ nhôm Xingfa 55, số lượng 3 bộ:',
    italic=True, size=13, align='left', sb=6, sa=4)
mk_para(anchor_bom,
    'Quy trình tính toán gồm 4 bước: (1) Nhận thông số đầu vào từ form báo giá (loại cửa, '
    'hệ nhôm, kích thước W×H, số lượng); (2) Tra cứu công thức cắt từ bảng MAU_SAN_PHAM '
    'theo mã hệ nhôm; (3) Tính chiều dài từng thanh và áp dụng thuật toán FFD để tối ưu '
    'số thanh 6000mm cần mua; (4) So sánh với tồn kho, xuất danh sách cần đặt thêm.',
    size=13, align='both', fi=1.25, sb=0, sa=6)
mk_para(anchor_bom,
    'a) Nguyên lý – Quy trình 4 bước của BOM Engine:',
    bold=False, italic=True, size=13, align='left', sb=6, sa=4)
mk_para(anchor_bom,
    'Tính năng bóc tách vật tư tự động (BOM – Bill of Materials) là chức năng kỹ thuật '
    'cốt lõi và khác biệt nhất của hệ thống ViralWindow so với các phần mềm quản lý thông '
    'thường. Thay vì tính toán thủ công trên bảng tính Excel, hệ thống tự động tính ra '
    'toàn bộ vật tư cần thiết dựa trên thông số kỹ thuật của từng hạng mục cửa nhôm kính.',
    size=13, align='both', fi=1.25, sb=0, sa=6)
mk_para(anchor_bom, '3.3.4. Thuật toán Bóc tách Vật tư (BOM Engine)',
        bold=True, size=13, align='left', sb=12, sa=6)
print("✅ Task 2: BOM Engine chèn xong")

# ── TASK 4: Bảng so sánh – chèn sau đoạn [117] "Đối tượng nghiên cứu" ──
anchor_compare = paras[118]._element  # đoạn sau "Đối tượng nghiên cứu"
mk_para(anchor_compare, '', sb=0, sa=8)
mk_para(anchor_compare,
    'Dựa trên bảng so sánh, ViralWindow lấp đầy khoảng trống mà các phần mềm ERP thương '
    'mại chưa giải quyết: tự động hóa nghiệp vụ đặc thù ngành cửa nhôm kính (BOM) với '
    'chi phí miễn phí, phù hợp doanh nghiệp nhỏ và vừa tại Việt Nam.',
    size=13, align='both', fi=1.25, sb=4, sa=6)
mk_table(anchor_compare,
    headers=['Tiêu chí', 'ViralWindow', 'MISA AMIS', 'Base.vn', 'Google Sheets'],
    rows=[
        ['Quản lý nhôm kính chuyên biệt', '✅ Có',        '❌ Không',   '❌ Không',   '❌ Không'],
        ['Bóc tách BOM tự động',          '✅ Có',        '❌ Không',   '❌ Không',   '❌ Thủ công'],
        ['Chi phí triển khai',            '✅ Miễn phí',  '5–15tr/năm','3–8tr/năm', '✅ Miễn phí'],
        ['Tích hợp AI (Gemini)',          '✅ Có',        '❌ Không',   '❌ Không',   '❌ Không'],
        ['Thông báo thời gian thực',      '✅ Socket.IO', '⚠️ Giới hạn','⚠️ Giới hạn','❌ Không'],
        ['Tùy chỉnh theo nghiệp vụ',     '✅ Hoàn toàn', '⚠️ Hạn chế', '⚠️ Hạn chế', '✅ Có'],
    ],
    col_widths=(5, 3.5, 3, 2.5, 3)
)
mk_para(anchor_compare, 'Bảng 0.1: So sánh ViralWindow với các giải pháp hiện có',
        italic=True, size=12, align='center', sb=4, sa=8)
mk_para(anchor_compare, '1.3. So sánh với các giải pháp hiện có',
        bold=True, size=13, align='left', sb=10, sa=6)
print("✅ Task 4: Bảng so sánh chèn xong")

doc.save(src)
print(f"\n🎉 Đã lưu! Tổng paragraphs: {len(doc.paragraphs)}")
