# -*- coding: utf-8 -*-
"""Xác minh nội dung DANH MỤC TỪ VIẾT TẮT trong FINAL.docx"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"
doc = Document(doc_path)

in_abbr = False
abbr_lines = []
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if "DANH MỤC TỪ VIẾT TẮT" in txt.upper():
        in_abbr = True
        abbr_lines.append(f"[{i}] TIÊU ĐỀ: {txt}")
        continue
    if in_abbr:
        if "MỞ ĐẦU" in txt.upper() and len(txt) < 20:
            abbr_lines.append(f"[{i}] >>> KẾT THÚC – MỞ ĐẦU <<<")
            break
        if txt:
            abbr_lines.append(f"[{i}] {txt[:90]}")

print(f"Tổng đoạn: {len(doc.paragraphs)}")
print(f"Số dòng trong phần từ viết tắt: {len(abbr_lines)}")
print()
for line in abbr_lines:
    print(line)
