# -*- coding: utf-8 -*-
"""Xác minh DANH MỤC TỪ VIẾT TẮT đã được chèn đúng vị trí"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"
doc = Document(doc_path)

found_abbr = False
found_mo_dau = False
context = []

for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if "DANH MỤC TỪ VIẾT TẮT" in txt.upper():
        found_abbr = True
        context.append(f"[{i}] >>> DANH MỤC TỪ VIẾT TẮT <<<")
    elif "MỞ ĐẦU" in txt.upper() and len(txt) < 20:
        found_mo_dau = True
        context.append(f"[{i}] >>> MỞ ĐẦU <<<")
    elif found_abbr and not found_mo_dau and txt:
        context.append(f"[{i}] {txt[:80]}")

print("=== XÁC MINH VỊ TRÍ ===")
for line in context[:40]:
    print(line)

print(f"\n✅ Có DANH MỤC TỪ VIẾT TẮT: {found_abbr}")
print(f"✅ Có MỞ ĐẦU sau: {found_mo_dau}")
print(f"✅ Tổng đoạn trong tài liệu: {len(doc.paragraphs)}")
