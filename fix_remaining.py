# -*- coding: utf-8 -*-
"""Sửa 2 lỗi nhỏ còn sót: 3.3.2 Trang chủ và 3.4.3 Cài đặt"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"
doc = Document(path)

fixes = {
    '3.3.2. Giao diện Trang chủ':            '3.4.2. Giao diện Trang chủ',
    '3.4.3. Giao diện Cài đặt Hệ thống':    '3.5.3. Giao diện Cài đặt Hệ thống',
    '3.4.5. Giao diện Quản lý Kho vật tư vật tư': '3.4.5. Giao diện Quản lý Kho vật tư',
}

changed = 0
for para in doc.paragraphs:
    for old, new in fixes.items():
        if old in para.text:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    print(f"Fixed: {old} -> {new}")
                    changed += 1

doc.save(path)
print(f"\nDone: fixed {changed} items")

# Xac nhan
doc2 = Document(path)
print("\nMuc Chuong 3 sau sua:")
for p in doc2.paragraphs:
    t = p.text.strip()
    if t.startswith('3.') and 'Giao dien' in t or ('3.' in t and len(t)<80 and t[0]=='3'):
        if len(t) < 80:
            print(f"  {t}")
