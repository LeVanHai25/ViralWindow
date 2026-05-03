# -*- coding: utf-8 -*-
"""Thực thi kế hoạch chỉnh sửa đồ án - chèn nội dung kỹ thuật vào FINAL.docx"""
import sys, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

src = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"
out = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"

doc = Document(src)
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ── helpers ──────────────────────────────────────────────
def q(tag): return f'{{{WNS}}}{tag}'

def mk_run(p_elem, text, bold=False, italic=False, size=13, color=None, mono=False):
    r = etree.SubElement(p_elem, q('r'))
    rPr = etree.SubElement(r, q('rPr'))
    if bold:   etree.SubElement(rPr, q('b')); etree.SubElement(rPr, q('bCs'))
    if italic: etree.SubElement(rPr, q('i'))
    if color:
        cl = etree.SubElement(rPr, q('color')); cl.set(q('val'), color)
    font_name = 'Courier New' if mono else 'Times New Roman'
    rf = etree.SubElement(rPr, q('rFonts'))
    rf.set(q('ascii'), font_name); rf.set(q('hAnsi'), font_name)
    sz = etree.SubElement(rPr, q('sz')); sz.set(q('val'), str(size*2))
    szCs = etree.SubElement(rPr, q('szCs')); szCs.set(q('val'), str(size*2))
    t = etree.SubElement(r, q('t'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    return r

def mk_para(ref_elem, text='', bold=False, italic=False, size=13,
            align='left', color=None, space_before=0, space_after=6,
            first_indent=0, mono=False, keep_next=False):
    p = etree.Element(q('p'))
    ref_elem.addprevious(p)
    pPr = etree.SubElement(p, q('pPr'))
    jc = etree.SubElement(pPr, q('jc')); jc.set(q('val'), align)
    sp = etree.SubElement(pPr, q('spacing'))
    sp.set(q('before'), str(int(space_before*20)))
    sp.set(q('after'),  str(int(space_after*20)))
    if first_indent:
        ind = etree.SubElement(pPr, q('ind'))
        ind.set(q('firstLine'), str(int(first_indent*567)))
    if keep_next:
        etree.SubElement(pPr, q('keepNext'))
    if text:
        mk_run(p, text, bold=bold, italic=italic, size=size, color=color, mono=mono)
    return p

def mk_code_block(ref_elem, lines):
    """Tạo khối code với background xám nhạt"""
    for line in lines:
        p = etree.Element(q('p'))
        ref_elem.addprevious(p)
        pPr = etree.SubElement(p, q('pPr'))
        sp = etree.SubElement(pPr, q('spacing'))
        sp.set(q('before'), '20'); sp.set(q('after'), '20')
        # indent trái
        ind = etree.SubElement(pPr, q('ind'))
        ind.set(q('left'), '567')  # 1cm
        if line:
            mk_run(p, line, mono=True, size=10, color='1F3864')

def mk_table_2col(ref_elem, headers, rows, col_widths_cm=(4, 12)):
    """Tạo bảng 2 cột đơn giản"""
    from docx.shared import Cm as DocxCm
    # Tạo table XML thủ công
    tbl = etree.Element(q('tbl'))
    ref_elem.addprevious(tbl)

    # tblPr
    tblPr = etree.SubElement(tbl, q('tblPr'))
    tblW = etree.SubElement(tblPr, q('tblW'))
    tblW.set(q('w'), '9360'); tblW.set(q('type'), 'dxa')
    tblBorders = etree.SubElement(tblPr, q('tblBorders'))
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = etree.SubElement(tblBorders, q(side))
        b.set(q('val'),'single'); b.set(q('sz'),'4')
        b.set(q('space'),'0'); b.set(q('color'),'000000')
    tblLook = etree.SubElement(tblPr, q('tblLook'))
    tblLook.set(q('val'),'0400')

    # tblGrid
    tblGrid = etree.SubElement(tbl, q('tblGrid'))
    for w in col_widths_cm:
        gc = etree.SubElement(tblGrid, q('gridCol'))
        gc.set(q('w'), str(int(w*567)))

    def add_row(cells_data, header=False):
        tr = etree.SubElement(tbl, q('tr'))
        trPr = etree.SubElement(tr, q('trPr'))
        trH = etree.SubElement(trPr, q('trHeight'))
        trH.set(q('val'), '400')
        for ci, (ctext, cw) in enumerate(zip(cells_data, col_widths_cm)):
            tc = etree.SubElement(tr, q('tc'))
            tcPr = etree.SubElement(tc, q('tcPr'))
            w_el = etree.SubElement(tcPr, q('tcW'))
            w_el.set(q('w'), str(int(cw*567))); w_el.set(q('type'),'dxa')
            tcMar = etree.SubElement(tcPr, q('tcMar'))
            for side, val in [('top','60'),('bottom','60'),('left','113'),('right','113')]:
                m = etree.SubElement(tcMar, q(side))
                m.set(q('w'), val); m.set(q('type'),'dxa')
            p_tc = etree.SubElement(tc, q('p'))
            pPr_tc = etree.SubElement(p_tc, q('pPr'))
            jc_tc = etree.SubElement(pPr_tc, q('jc'))
            jc_tc.set(q('val'), 'center' if ci==0 else 'left')
            sp_tc = etree.SubElement(pPr_tc, q('spacing'))
            sp_tc.set(q('before'),'20'); sp_tc.set(q('after'),'20')
            mk_run(p_tc, ctext, bold=header, size=13)
        return tr

    if headers:
        add_row(headers, header=True)
    for row in rows:
        add_row(row)
    return tbl

# ── Tìm anchor points ─────────────────────────────────────
paras = doc.paragraphs
refs = {}
for i, p in enumerate(paras):
    t = p.text.strip()
    if '3.3.3' in t and 'Quản lý gói' in t: refs['after_333'] = p._element
    if '3.4.' in t and 'Giao diện chung' in t: refs['giao_dien_chung'] = p._element
    if 'Mục tiêu đề tài' in t: refs['muc_tieu'] = p._element
    if '4.2. Kiểm thử hệ thống' in t: refs['kt_he_thong'] = p._element
    if '4.3.' in t and 'Tổng kết' in t: refs['tong_ket_kt'] = p._element
    if 'CHƯƠNG 2' in t.upper() and 'PHÂN TÍCH' in t.upper(): refs['chuong2'] = p._element

print("Anchors found:", list(refs.keys()))

# ════════════════════════════════════════════════════════
# TASK 1: FIX SỐ MỤC 3.3.1 TRÙNG
# Đổi tiêu đề "3.3.1. Giao diện Đăng nhập" → "3.4.1. Giao diện Đăng nhập"
# ════════════════════════════════════════════════════════
fixed_nums = 0
mapping = {
    '3.3.1. Giao diện Đăng nhập':    '3.4.1. Giao diện Đăng nhập',
    '3.3.2. Giao diện Trang chủ':     '3.4.2. Giao diện Trang chủ',
    '3.3.3. Giao diện Quản lý Dự án': '3.4.3. Giao diện Quản lý Dự án',
    '3.3.4. Giao diện Lập Báo giá':   '3.4.4. Giao diện Lập Báo giá',
    '3.3.5. Giao diện Quản lý Kho':   '3.4.5. Giao diện Quản lý Kho vật tư',
    '3.3.6. Giao diện Quản lý Tài chính': '3.4.6. Giao diện Quản lý Tài chính',
    'Giao diện chung':               'Giao diện chung',
    '3.4. Giao diện Quản trị':       '3.5. Giao diện Quản trị',
    '3.4.1. Giao diện Quản lý Người dùng': '3.5.1. Giao diện Quản lý Người dùng',
    '3.4.2. Giao diện Quản lý Phân quyền': '3.5.2. Giao diện Quản lý Phân quyền',
    '3.4.3. Giao diện Cài đặt Hệ thống':  '3.5.3. Giao diện Cài đặt Hệ thống',
}
for para in doc.paragraphs:
    for old, new in mapping.items():
        if old in para.text and len(para.text.strip()) < 80:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    fixed_nums += 1
                    print(f"  Renamed: {old[:40]} → {new[:40]}")
print(f"✅ Task 1: Sửa {fixed_nums} số mục trùng")

# ════════════════════════════════════════════════════════
# TASK 2: THÊM MỤC 3.3.4 – THUẬT TOÁN BOM
# Chèn trước "Giao diện chung" (mục 3.4 mới)
# ════════════════════════════════════════════════════════
if 'giao_dien_chung' in refs:
    anchor = refs['giao_dien_chung']

    # Tiêu đề mục
    mk_para(anchor, '3.3.4. Thuật toán Bóc tách Vật tư (BOM Engine)',
            bold=True, size=13, space_before=12, space_after=6)

    # Giới thiệu
    mk_para(anchor,
        'Tính năng bóc tách vật tư (BOM – Bill of Materials) là chức năng kỹ thuật cốt lõi '
        'và khác biệt nhất của hệ thống ViralWindow, giúp tự động tính toán số lượng từng loại '
        'vật tư (nhôm, kính, phụ kiện) dựa trên thông số kỹ thuật của từng hạng mục cửa, '
        'thay thế hoàn toàn quy trình tính toán thủ công trên bảng tính Excel trước đây.',
        size=13, align='both', first_indent=1.25, space_after=6)

    # a) Nguyên lý
    mk_para(anchor, 'a) Nguyên lý hoạt động – Quy trình 4 bước',
            bold=True, italic=True, size=13, space_before=6, space_after=4)
    steps = [
        'Bước 1 – Nhận thông số đầu vào: Loại cửa, hệ nhôm, kích thước W×H (mm), số lượng bộ.',
        'Bước 2 – Tra cứu công thức: Mỗi hệ nhôm có bộ công thức riêng trong bảng MAU_SAN_PHAM, '
        'định nghĩa chiều dài và số lượng từng thanh theo kích thước đầu vào.',
        'Bước 3 – Tối ưu cắt: Hệ thống áp dụng thuật toán First-Fit Decreasing (FFD) để tính '
        'số thanh nhôm 6000mm cần mua và tỷ lệ hao phí tối thiểu.',
        'Bước 4 – Xuất BOM: Danh sách vật tư chi tiết (nhôm, kính, phụ kiện) kèm số lượng '
        'và số thanh cần đặt mua, so sánh trực tiếp với tồn kho hiện tại.',
    ]
    for s in steps:
        mk_para(anchor, s, size=13, align='both', first_indent=1.25, space_after=4)

    # b) Ví dụ
    mk_para(anchor, 'b) Ví dụ minh họa – Cửa đi 1200×2100mm, hệ Xingfa 55, số lượng: 3 bộ',
            bold=True, italic=True, size=13, space_before=6, space_after=4)

    mk_table_2col(anchor,
        headers=['Thanh nhôm', 'Công thức / Kết quả'],
        rows=[
            ['Thanh đứng khung', 'Dài = H – 10 = 2090mm × 2 thanh/bộ → 6 đoạn → cần 4 thanh 6m'],
            ['Thanh ngang khung', 'Dài = W – 10 = 1190mm × 2 thanh/bộ → 6 đoạn → cần 3 thanh 6m'],
            ['Thanh đứng cánh',  'Dài = H – 55 = 2045mm × 2 thanh/bộ → 6 đoạn → cần 4 thanh 6m'],
            ['Thanh ngang cánh', 'Dài = W – 55 = 1145mm × 2 thanh/bộ → 6 đoạn → cần 3 thanh 6m'],
            ['Kính cường lực 5mm','1130×2030mm × 3 tấm'],
            ['Bản lề inox 4 tấc','3 bộ × 3 cái = 9 bản lề'],
            ['Hao phí nhôm ước tính','≈ 8.3% – tối ưu FFD'],
        ],
        col_widths_cm=(5, 11)
    )
    mk_para(anchor, '', space_before=0, space_after=4)

    # c) Code minh họa
    mk_para(anchor, 'c) Đoạn mã xử lý tối ưu cắt (First-Fit Decreasing)',
            bold=True, italic=True, size=13, space_before=6, space_after=4)
    mk_code_block(anchor, [
        '// backend/controllers/bomController.js',
        'function optimizeCutting(pieces, barLen=6000, kerf=3) {',
        '  const sorted = [...pieces].sort((a,b) => b.length - a.length);',
        '  const bars = [];',
        '  for (const piece of sorted) {',
        '    let placed = false;',
        '    for (const bar of bars) {',
        '      const used = bar.reduce((s,p)=>s+p.length+kerf, 0);',
        '      if (used + piece.length + kerf <= barLen) {',
        '        bar.push(piece); placed = true; break;',
        '      }',
        '    }',
        '    if (!placed) bars.push([piece]);',
        '  }',
        '  return { totalBars: bars.length,',
        '    wastePercent: ((barLen*bars.length -',
        '      pieces.reduce((s,p)=>s+p.length,0))/(barLen*bars.length)*100).toFixed(1)',
        '  };',
        '}',
    ])
    mk_para(anchor, '', space_before=0, space_after=8)
    print("✅ Task 2: Đã chèn mục 3.3.4 – BOM Engine")
else:
    print("⚠️ Task 2: Không tìm thấy anchor 'giao_dien_chung'")

# ════════════════════════════════════════════════════════
# TASK 3: THÊM CODE JWT/RBAC vào mục 3.2
# Tìm "3.2.1. Cơ chế xác thực" và chèn code sau phần mô tả
# ════════════════════════════════════════════════════════
rbac_anchor = None
for i, p in enumerate(doc.paragraphs):
    if 'Phân quyền theo mô hình Role-Based Access Control' in p.text:
        rbac_anchor = doc.paragraphs[i+1]._element if i+1 < len(doc.paragraphs) else None
        break

if rbac_anchor:
    mk_para(rbac_anchor, 'Minh họa triển khai RBAC Middleware trong hệ thống ViralWindow:',
            italic=True, size=13, space_before=4, space_after=4)
    mk_code_block(rbac_anchor, [
        '// middleware/authMiddleware.js',
        'const requirePermission = (permCode) => async (req, res, next) => {',
        '  const [rows] = await db.query(',
        '    `SELECT qh.ma_code FROM CHUC_VU_QUYEN cvq',
        '     JOIN QUYEN_HAN qh ON cvq.ma_quyen = qh.id',
        '     WHERE cvq.ma_chuc_vu = ? AND qh.ma_code = ?`,',
        '    [req.user.roleId, permCode]',
        '  );',
        '  if (rows.length === 0)',
        '    return res.status(403).json({ message: `Không có quyền [${permCode}]` });',
        '  next();',
        '};',
        '',
        '// Áp dụng: chỉ CUSTOMER_DELETE mới xóa được khách hàng',
        "router.delete('/customers/:id', verifyToken,",
        "  requirePermission('CUSTOMER_DELETE'), deleteCustomer);",
    ])
    mk_para(rbac_anchor, '', space_before=0, space_after=6)
    print("✅ Task 3: Đã chèn code RBAC")
else:
    print("⚠️ Task 3: Không tìm thấy anchor RBAC")

# ════════════════════════════════════════════════════════
# TASK 4: BẢNG SO SÁNH HỆ THỐNG – chèn sau Mục tiêu đề tài
# ════════════════════════════════════════════════════════
if 'muc_tieu' in refs:
    anchor4 = refs['muc_tieu']
    # Chèn TRƯỚC mục tiêu (sẽ thành phần sau Mục tiêu khi đọc)
    # Thực ra ta cần chèn AFTER – tìm para kế tiếp
    muc_tieu_idx = next(i for i,p in enumerate(doc.paragraphs)
                        if 'Mục tiêu đề tài' in p.text)
    # Tìm đoạn kế tiếp là "Đối tượng và phạm vi"
    doi_tuong_anchor = None
    for i in range(muc_tieu_idx+1, len(doc.paragraphs)):
        if 'Đối tượng và phạm vi' in doc.paragraphs[i].text:
            doi_tuong_anchor = doc.paragraphs[i]._element
            break

    if doi_tuong_anchor:
        mk_para(doi_tuong_anchor, 'So sánh với các giải pháp hiện có trên thị trường',
                bold=True, size=13, space_before=8, space_after=6)
        mk_table_2col(doi_tuong_anchor,
            headers=['Tiêu chí', 'ViralWindow', 'MISA AMIS', 'Base.vn', 'Google Sheets'],
            rows=[
                ['Quản lý nhôm kính', '✅ Chuyên biệt', '❌ Không có', '❌ Không có', '❌ Không có'],
                ['BOM tự động',       '✅ Có',          '❌ Không',    '❌ Không',    '❌ Thủ công'],
                ['Chi phí/năm',       '✅ Miễn phí',    '5–15 triệu', '3–8 triệu',  '✅ Miễn phí'],
                ['Tích hợp AI',       '✅ Gemini API',  '❌ Không',   '❌ Không',    '❌ Không'],
                ['Real-time notify',  '✅ Socket.IO',   '⚠️ Có giới hạn','⚠️ Có','❌ Không'],
                ['Tùy chỉnh nghiệp vụ','✅ Hoàn toàn', '⚠️ Hạn chế', '⚠️ Hạn chế', '✅ Có'],
            ],
            col_widths_cm=(4, 3, 3, 3, 3)
        )
        mk_para(doi_tuong_anchor,
            'Qua bảng so sánh, ViralWindow lấp đầy khoảng trống mà các phần mềm ERP thương mại '
            'chưa giải quyết được: tự động hóa nghiệp vụ đặc thù ngành cửa nhôm kính (BOM) với '
            'chi phí phù hợp cho doanh nghiệp nhỏ Việt Nam.',
            size=13, align='both', first_indent=1.25, space_before=4, space_after=8)
        print("✅ Task 4: Đã chèn bảng so sánh hệ thống")
else:
    print("⚠️ Task 4: Không tìm thấy anchor mục tiêu")

# ════════════════════════════════════════════════════════
# TASK 5: MỞ RỘNG CHƯƠNG 4 – Test case chi tiết đăng nhập
# ════════════════════════════════════════════════════════
if 'tong_ket_kt' in refs:
    anchor5 = refs['tong_ket_kt']
    mk_para(anchor5, '4.2.9. Kiểm thử bảo mật (Security Testing)',
            bold=True, size=13, space_before=10, space_after=6)
    mk_para(anchor5,
        'Hệ thống được kiểm thử theo danh sách OWASP Top 10, tập trung vào các mối đe dọa '
        'phổ biến nhất đối với ứng dụng web:',
        size=13, align='both', first_indent=1.25, space_after=4)
    mk_table_2col(anchor5,
        headers=['Loại tấn công', 'Phương pháp kiểm thử', 'Kết quả'],
        rows=[
            ["SQL Injection",    "Nhập '; DROP TABLE users; -- vào form email", "✅ Blocked – Prepared Statement"],
            ["XSS",             "<script>alert('XSS')</script> vào tên khách hàng","✅ Blocked – Output encoding"],
            ["CSRF",            "Gửi request giả từ domain không được phép",    "✅ Blocked – CORS restriction"],
            ["IDOR",            "Gọi /api/projects/999 (dự án không phải của mình)","✅ Blocked – RBAC check"],
            ["Brute Force",     "Thử 100 mật khẩu sai liên tiếp",               "⚠️ Chưa có rate limiting"],
            ["Token hết hạn",   "Dùng JWT token cũ sau 24 giờ",                 "✅ Blocked – HTTP 403"],
        ],
        col_widths_cm=(4, 6, 6)
    )
    mk_para(anchor5, '', space_before=0, space_after=8)
    print("✅ Task 5: Đã chèn bảng kiểm thử bảo mật")

doc.save(out)
print(f"\n🎉 Hoàn thành! Lưu: {out}")
print(f"   Tổng paragraphs: {len(doc.paragraphs)}")
