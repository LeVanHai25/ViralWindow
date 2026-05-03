# -*- coding: utf-8 -*-
"""
Chèn DANH MỤC TỪ VIẾT TẮT bằng python-docx API chuẩn
Dùng cách chèn vào body.xml trực tiếp - đáng tin cậy hơn
"""
import sys, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

doc_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_v2.docx"
out_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"

ABBREVIATIONS = sorted([
    ("2FA",     "Two-Factor Authentication",                    "Xác thực hai yếu tố"),
    ("API",     "Application Programming Interface",            "Giao diện lập trình ứng dụng"),
    ("BOM",     "Bill of Materials",                            "Danh mục vật tư"),
    ("CORS",    "Cross-Origin Resource Sharing",                "Chia sẻ tài nguyên giữa các nguồn gốc"),
    ("CRUD",    "Create, Read, Update, Delete",                 "Tạo, Đọc, Cập nhật, Xóa"),
    ("CSRF",    "Cross-Site Request Forgery",                   "Tấn công giả mạo yêu cầu liên trang"),
    ("CSS",     "Cascading Style Sheets",                       "Ngôn ngữ định kiểu tầng"),
    ("DCL",     "Data Control Language",                        "Ngôn ngữ kiểm soát dữ liệu"),
    ("DDL",     "Data Definition Language",                     "Ngôn ngữ định nghĩa dữ liệu"),
    ("ERP",     "Enterprise Resource Planning",                 "Hệ thống hoạch định nguồn lực doanh nghiệp"),
    ("HTML",    "HyperText Markup Language",                    "Ngôn ngữ đánh dấu siêu văn bản"),
    ("HTTP",    "HyperText Transfer Protocol",                  "Giao thức truyền tải siêu văn bản"),
    ("HTTPS",   "HyperText Transfer Protocol Secure",           "Giao thức HTTP bảo mật (có mã hóa SSL/TLS)"),
    ("I/O",     "Input / Output",                               "Đầu vào / Đầu ra"),
    ("JWT",     "JSON Web Token",                               "Mã thông báo xác thực dạng JSON"),
    ("KPI",     "Key Performance Indicator",                    "Chỉ số hiệu suất chính"),
    ("MySQL",   "My Structured Query Language",                 "Hệ quản trị cơ sở dữ liệu quan hệ"),
    ("npm",     "Node Package Manager",                         "Trình quản lý gói của Node.js"),
    ("OWASP",   "Open Web Application Security Project",        "Dự án bảo mật ứng dụng web mở"),
    ("PDF",     "Portable Document Format",                     "Định dạng tài liệu di động"),
    ("RBAC",    "Role-Based Access Control",                    "Kiểm soát truy cập dựa trên vai trò"),
    ("RDBMS",   "Relational Database Management System",        "Hệ quản trị cơ sở dữ liệu quan hệ"),
    ("REST",    "Representational State Transfer",              "Kiến trúc truyền trạng thái đại diện"),
    ("RWD",     "Responsive Web Design",                        "Thiết kế web thích ứng"),
    ("SEO",     "Search Engine Optimization",                   "Tối ưu hóa công cụ tìm kiếm"),
    ("SQL",     "Structured Query Language",                    "Ngôn ngữ truy vấn có cấu trúc"),
    ("SSL",     "Secure Sockets Layer",                         "Lớp kết nối bảo mật"),
    ("TCL",     "Transaction Control Language",                 "Ngôn ngữ kiểm soát giao dịch"),
    ("TLS",     "Transport Layer Security",                     "Bảo mật tầng vận chuyển"),
    ("UML",     "Unified Modeling Language",                    "Ngôn ngữ mô hình hóa thống nhất"),
    ("URL",     "Uniform Resource Locator",                     "Định vị tài nguyên thống nhất"),
    ("UX",      "User Experience",                              "Trải nghiệm người dùng"),
    ("XSS",     "Cross-Site Scripting",                         "Tấn công kịch bản liên trang"),
    ("XHTML",   "Extensible HyperText Markup Language",         "Ngôn ngữ HTML mở rộng"),
], key=lambda x: x[0].upper())

doc = Document(doc_path)

# --- Tìm vị trí MỞ ĐẦU ---
insert_before_idx = None
for i, para in enumerate(doc.paragraphs):
    if "MỞ ĐẦU" in para.text.strip().upper() and len(para.text.strip()) < 20:
        insert_before_idx = i
        break

if insert_before_idx is None:
    print("❌ Không tìm thấy MỞ ĐẦU")
    sys.exit(1)

print(f"✅ Sẽ chèn trước đoạn [{insert_before_idx}]: '{doc.paragraphs[insert_before_idx].text}'")

# --- Hàm tạo paragraph chuẩn ---
def new_para(doc, text="", bold=False, size=13, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=0, first_indent_cm=0, color_hex=None,
             italic=False, keep_with_next=False):
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p.paragraph_format.alignment     = align
    p.paragraph_format.space_before  = Pt(space_before)
    p.paragraph_format.space_after   = Pt(space_after)
    if first_indent_cm:
        p.paragraph_format.first_line_indent = Cm(first_indent_cm)
    if text:
        run = p.add_run(text)
        run.font.name   = "Times New Roman"
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        if color_hex:
            from docx.shared import RGBColor
            r, g, b = int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16)
            run.font.color.rgb = RGBColor(r, g, b)
    if keep_with_next:
        pPr = p._p.get_or_add_pPr()
        kwn = OxmlElement('w:keepNext')
        pPr.append(kwn)
    return p

# --- Hàm chèn paragraph vào vị trí cụ thể ---
def insert_para_before(doc, ref_para_idx, new_p):
    """Chèn paragraph p trước paragraph tại index ref_para_idx"""
    ref_para = doc.paragraphs[ref_para_idx]
    ref_para._element.addprevious(new_p._element)

# --- Tạo và chèn nội dung ---
# Ta sẽ add vào cuối rồi move, hoặc dùng addprevious trực tiếp
# Approach: tạo tất cả paragraphs, rồi chèn trước MO DAU

ref_elem = doc.paragraphs[insert_before_idx]._element

def insert_xml_before(ref_elem, text, bold=False, size=13,
                      align="left", color=None, italic=False,
                      space_before=0, space_after=0,
                      first_indent=0, page_break_before=False):
    """Tạo và chèn một paragraph XML element trước ref_elem"""
    WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    p = etree.SubElement(ref_elem.getparent(), f'{{{WNS}}}p')
    ref_elem.addprevious(p)
    
    pPr = etree.SubElement(p, f'{{{WNS}}}pPr')
    
    # Alignment
    jc = etree.SubElement(pPr, f'{{{WNS}}}jc')
    jc.set(f'{{{WNS}}}val', align)
    
    # Spacing
    sp = etree.SubElement(pPr, f'{{{WNS}}}spacing')
    sp.set(f'{{{WNS}}}before', str(int(space_before * 20)))
    sp.set(f'{{{WNS}}}after',  str(int(space_after * 20)))
    
    # First line indent
    if first_indent:
        ind = etree.SubElement(pPr, f'{{{WNS}}}ind')
        ind.set(f'{{{WNS}}}firstLine', str(int(first_indent * 567)))
    
    # Page break before
    if page_break_before:
        pbk = etree.SubElement(pPr, f'{{{WNS}}}pageBreakBefore')
    
    if not text:
        return p
    
    r = etree.SubElement(p, f'{{{WNS}}}r')
    rPr = etree.SubElement(r, f'{{{WNS}}}rPr')
    
    if bold:
        etree.SubElement(rPr, f'{{{WNS}}}b')
        etree.SubElement(rPr, f'{{{WNS}}}bCs')
    if italic:
        etree.SubElement(rPr, f'{{{WNS}}}i')
    if color:
        cl = etree.SubElement(rPr, f'{{{WNS}}}color')
        cl.set(f'{{{WNS}}}val', color)
    
    sz = etree.SubElement(rPr, f'{{{WNS}}}sz')
    sz.set(f'{{{WNS}}}val', str(size * 2))
    szCs = etree.SubElement(rPr, f'{{{WNS}}}szCs')
    szCs.set(f'{{{WNS}}}val', str(size * 2))
    
    rFonts = etree.SubElement(rPr, f'{{{WNS}}}rFonts')
    rFonts.set(f'{{{WNS}}}ascii',  'Times New Roman')
    rFonts.set(f'{{{WNS}}}hAnsi', 'Times New Roman')
    rFonts.set(f'{{{WNS}}}cs',    'Times New Roman')
    
    t = etree.SubElement(r, f'{{{WNS}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    
    return p

def insert_abbr_row_before(ref_elem, abbr, full_en, full_vi):
    """Chèn một dòng từ viết tắt: [abbr] [tab] [full_en] [tab] [full_vi]"""
    WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    p = etree.Element(f'{{{WNS}}}p')
    ref_elem.addprevious(p)
    
    pPr = etree.SubElement(p, f'{{{WNS}}}pPr')
    sp = etree.SubElement(pPr, f'{{{WNS}}}spacing')
    sp.set(f'{{{WNS}}}before', '30')
    sp.set(f'{{{WNS}}}after',  '30')
    
    # Tab stops
    tabs = etree.SubElement(pPr, f'{{{WNS}}}tabs')
    for pos_twips, val in [('2268', 'left'), ('8505', 'left')]:  # 4cm, 15cm
        tab_el = etree.SubElement(tabs, f'{{{WNS}}}tab')
        tab_el.set(f'{{{WNS}}}val', val)
        tab_el.set(f'{{{WNS}}}pos', pos_twips)
    
    def make_run(text, bold=False, color=None):
        r = etree.SubElement(p, f'{{{WNS}}}r')
        rPr = etree.SubElement(r, f'{{{WNS}}}rPr')
        if bold:
            etree.SubElement(rPr, f'{{{WNS}}}b')
            etree.SubElement(rPr, f'{{{WNS}}}bCs')
        if color:
            cl = etree.SubElement(rPr, f'{{{WNS}}}color')
            cl.set(f'{{{WNS}}}val', color)
        sz = etree.SubElement(rPr, f'{{{WNS}}}sz'); sz.set(f'{{{WNS}}}val', '26')
        szCs = etree.SubElement(rPr, f'{{{WNS}}}szCs'); szCs.set(f'{{{WNS}}}val', '26')
        rf = etree.SubElement(rPr, f'{{{WNS}}}rFonts')
        rf.set(f'{{{WNS}}}ascii', 'Times New Roman')
        rf.set(f'{{{WNS}}}hAnsi', 'Times New Roman')
        t = etree.SubElement(r, f'{{{WNS}}}t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        return r
    
    def make_tab():
        r = etree.SubElement(p, f'{{{WNS}}}r')
        rPr = etree.SubElement(r, f'{{{WNS}}}rPr')
        sz = etree.SubElement(rPr, f'{{{WNS}}}sz'); sz.set(f'{{{WNS}}}val', '26')
        szCs = etree.SubElement(rPr, f'{{{WNS}}}szCs'); szCs.set(f'{{{WNS}}}val', '26')
        etree.SubElement(r, f'{{{WNS}}}tab')
        return r
    
    make_run(abbr, bold=True, color='1F4E79')
    make_tab()
    make_run(full_en)
    make_tab()
    make_run(full_vi)

# ---- Chèn theo thứ tự NGƯỢC (vì addprevious chèn trước ref_elem liên tục) ----
# Cuối cùng đến đầu tiên

# 1. Dòng trống sau bảng + page break
insert_xml_before(ref_elem, "", space_before=0, space_after=0)

# 2. Chèn từng dòng (ngược lại)
for abbr, full_en, full_vi in reversed(ABBREVIATIONS):
    insert_abbr_row_before(ref_elem, abbr, full_en, full_vi)

# 3. Dòng header (in đậm, gạch chân)
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
p_hdr = etree.Element(f'{{{WNS}}}p')
ref_elem.addprevious(p_hdr)
pPr_h = etree.SubElement(p_hdr, f'{{{WNS}}}pPr')
sp_h = etree.SubElement(pPr_h, f'{{{WNS}}}spacing')
sp_h.set(f'{{{WNS}}}before', '60'); sp_h.set(f'{{{WNS}}}after', '60')
tabs_h = etree.SubElement(pPr_h, f'{{{WNS}}}tabs')
for pos_twips, val in [('2268', 'left'), ('8505', 'left')]:
    tab_el = etree.SubElement(tabs_h, f'{{{WNS}}}tab')
    tab_el.set(f'{{{WNS}}}val', val)
    tab_el.set(f'{{{WNS}}}pos', pos_twips)
pBdr_h = etree.SubElement(pPr_h, f'{{{WNS}}}pBdr')
bot = etree.SubElement(pBdr_h, f'{{{WNS}}}bottom')
bot.set(f'{{{WNS}}}val', 'single'); bot.set(f'{{{WNS}}}sz', '6')
bot.set(f'{{{WNS}}}space', '1'); bot.set(f'{{{WNS}}}color', '1F4E79')
r_hdr = etree.SubElement(p_hdr, f'{{{WNS}}}r')
rPr_h = etree.SubElement(r_hdr, f'{{{WNS}}}rPr')
etree.SubElement(rPr_h, f'{{{WNS}}}b'); etree.SubElement(rPr_h, f'{{{WNS}}}bCs')
etree.SubElement(rPr_h, f'{{{WNS}}}u').set(f'{{{WNS}}}val', 'single')
sz_h = etree.SubElement(rPr_h, f'{{{WNS}}}sz'); sz_h.set(f'{{{WNS}}}val', '26')
szCs_h = etree.SubElement(rPr_h, f'{{{WNS}}}szCs'); szCs_h.set(f'{{{WNS}}}val', '26')
rf_h = etree.SubElement(rPr_h, f'{{{WNS}}}rFonts')
rf_h.set(f'{{{WNS}}}ascii', 'Times New Roman'); rf_h.set(f'{{{WNS}}}hAnsi', 'Times New Roman')
t_hdr = etree.SubElement(r_hdr, f'{{{WNS}}}t')
t_hdr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
t_hdr.text = 'Từ viết tắt\t\tDiễn giải tiếng Anh\t\tÝ nghĩa tiếng Việt'

# 4. Dòng trống trước bảng
insert_xml_before(ref_elem, "", space_before=0, space_after=0)

# 5. Tiêu đề DANH MỤC TỪ VIẾT TẮT
insert_xml_before(ref_elem, "DANH MỤC TỪ VIẾT TẮT",
                  bold=True, size=13, align="center",
                  space_before=12, space_after=6,
                  page_break_before=False)

# 6. Ngắt trang từ trang trước (page break trước mục này)
insert_xml_before(ref_elem, "", page_break_before=True, space_before=0, space_after=0)

doc.save(out_path)
print(f"✅ Đã chèn {len(ABBREVIATIONS)} từ viết tắt")
print(f"✅ Đã lưu: {out_path}")
