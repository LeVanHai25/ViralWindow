# -*- coding: utf-8 -*-
"""
Chèn DANH MỤC TỪ VIẾT TẮT - PHƯƠNG PHÁP CHUẨN
Dùng python-docx Document.add_paragraph() hoàn toàn, 
rồi dùng lxml để move các paragraphs vào đúng vị trí
"""
import sys, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Dùng file v2 (trước khi có các lần chèn bị lỗi)
src_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_v2.docx"
out_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"

ABBREVIATIONS = sorted([
    ("2FA",    "Two-Factor Authentication",                   "Xác thực hai yếu tố"),
    ("API",    "Application Programming Interface",           "Giao diện lập trình ứng dụng"),
    ("BOM",    "Bill of Materials",                           "Danh mục vật tư"),
    ("CORS",   "Cross-Origin Resource Sharing",               "Chia sẻ tài nguyên giữa các nguồn gốc"),
    ("CRUD",   "Create, Read, Update, Delete",                "Tạo, Đọc, Cập nhật, Xóa"),
    ("CSRF",   "Cross-Site Request Forgery",                  "Giả mạo yêu cầu liên trang"),
    ("CSS",    "Cascading Style Sheets",                      "Ngôn ngữ định kiểu tầng"),
    ("DCL",    "Data Control Language",                       "Ngôn ngữ kiểm soát dữ liệu"),
    ("DDL",    "Data Definition Language",                    "Ngôn ngữ định nghĩa dữ liệu"),
    ("ERP",    "Enterprise Resource Planning",                "Hoạch định nguồn lực doanh nghiệp"),
    ("HTML",   "HyperText Markup Language",                   "Ngôn ngữ đánh dấu siêu văn bản"),
    ("HTTP",   "HyperText Transfer Protocol",                 "Giao thức truyền tải siêu văn bản"),
    ("HTTPS",  "HyperText Transfer Protocol Secure",          "Giao thức HTTP bảo mật"),
    ("I/O",    "Input / Output",                              "Đầu vào / Đầu ra"),
    ("JWT",    "JSON Web Token",                              "Mã thông báo xác thực JSON"),
    ("KPI",    "Key Performance Indicator",                   "Chỉ số hiệu suất chính"),
    ("MySQL",  "My Structured Query Language",                "Hệ quản trị CSDL quan hệ mã nguồn mở"),
    ("npm",    "Node Package Manager",                        "Trình quản lý gói của Node.js"),
    ("OWASP",  "Open Web Application Security Project",       "Dự án bảo mật ứng dụng web mở"),
    ("PDF",    "Portable Document Format",                    "Định dạng tài liệu di động"),
    ("RBAC",   "Role-Based Access Control",                   "Kiểm soát truy cập dựa trên vai trò"),
    ("RDBMS",  "Relational Database Management System",       "Hệ quản trị cơ sở dữ liệu quan hệ"),
    ("REST",   "Representational State Transfer",             "Kiến trúc truyền trạng thái đại diện"),
    ("RWD",    "Responsive Web Design",                       "Thiết kế web thích ứng"),
    ("SEO",    "Search Engine Optimization",                  "Tối ưu hóa công cụ tìm kiếm"),
    ("SQL",    "Structured Query Language",                   "Ngôn ngữ truy vấn có cấu trúc"),
    ("SSL",    "Secure Sockets Layer",                        "Lớp kết nối bảo mật"),
    ("TCL",    "Transaction Control Language",                "Ngôn ngữ kiểm soát giao dịch"),
    ("TLS",    "Transport Layer Security",                    "Bảo mật tầng vận chuyển"),
    ("UML",    "Unified Modeling Language",                   "Ngôn ngữ mô hình hóa thống nhất"),
    ("URL",    "Uniform Resource Locator",                    "Định vị tài nguyên thống nhất"),
    ("UX",     "User Experience",                             "Trải nghiệm người dùng"),
    ("XHTML",  "Extensible HyperText Markup Language",        "Ngôn ngữ HTML mở rộng"),
    ("XSS",    "Cross-Site Scripting",                        "Tấn công kịch bản liên trang"),
], key=lambda x: x[0].upper())

# ---- Load tài liệu gốc ----
doc = Document(src_path)

# ---- Cũng áp dụng lại fix lề và heading từ trước ----
for section in doc.sections:
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(3.0)

# ---- Tìm index của MỞ ĐẦU ----
mo_dau_idx = None
for i, para in enumerate(doc.paragraphs):
    if "MỞ ĐẦU" in para.text.strip().upper() and len(para.text.strip()) < 20:
        mo_dau_idx = i
        break

print(f"MỞ ĐẦU tại index: {mo_dau_idx} — '{doc.paragraphs[mo_dau_idx].text}'")

# ---- Lấy element tham chiếu ----
mo_dau_elem = doc.paragraphs[mo_dau_idx]._element
body = mo_dau_elem.getparent()

# ---- Hàm tiện ích tạo paragraph element bằng python-docx rồi detach ----
def build_para(text="", bold=False, italic=False, size=13,
               align=WD_ALIGN_PARAGRAPH.LEFT,
               color=None, space_before=0, space_after=4,
               first_indent=0, underline=False):
    """Tạo một paragraph (không gắn vào doc), trả về _element"""
    tmp = Document()  # tài liệu tạm
    p = tmp.add_paragraph()
    pf = p.paragraph_format
    pf.alignment    = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if first_indent:
        pf.first_line_indent = Cm(first_indent)
    if text:
        run = p.add_run(text)
        run.font.name      = "Times New Roman"
        run.font.size      = Pt(size)
        run.font.bold      = bold
        run.font.italic    = italic
        run.font.underline = underline
        if color:
            r, g, b = color
            run.font.color.rgb = RGBColor(r, g, b)
    elem = copy.deepcopy(p._element)
    return elem

def build_abbr_row(abbr, full_en, full_vi):
    """Tạo paragraph dạng: ABBR  |  Full English  |  Ý nghĩa VN"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    tmp = Document()
    p = tmp.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    
    # Tab stops: 3cm cho cột 2, 9.5cm cho cột 3
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    for pos_cm, leader in [(3.0, 'none'), (9.5, 'none')]:
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'left')
        tab.set(qn('w:leader'), leader)
        tab.set(qn('w:pos'), str(int(pos_cm * 567)))
        tabs.append(tab)
    pPr.insert(0, tabs)
    
    # Cột 1: Từ viết tắt (in đậm, màu xanh navy)
    r1 = p.add_run(abbr)
    r1.font.name  = "Times New Roman"; r1.font.size = Pt(13)
    r1.font.bold  = True
    r1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    # Tab
    p.add_run('\t')
    
    # Cột 2: Tiếng Anh
    r2 = p.add_run(full_en)
    r2.font.name  = "Times New Roman"; r2.font.size = Pt(13)
    r2.font.italic = True
    
    # Tab
    p.add_run('\t')
    
    # Cột 3: Tiếng Việt
    r3 = p.add_run(full_vi)
    r3.font.name  = "Times New Roman"; r3.font.size = Pt(13)
    
    return copy.deepcopy(p._element)

# ---- Danh sách các element cần chèn (theo thứ tự xuất hiện) ----
to_insert = []

# Page break trước mục
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
pb_elem = OxmlElement('w:p')
pb_pPr  = OxmlElement('w:pPr')
pb_r    = OxmlElement('w:r')
pb_br   = OxmlElement('w:br')
pb_br.set(qn('w:type'), 'page')
pb_r.append(pb_br)
pb_elem.append(pb_pPr)
pb_elem.append(pb_r)
to_insert.append(pb_elem)

# Tiêu đề
to_insert.append(build_para(
    "DANH MỤC TỪ VIẾT TẮT",
    bold=True, size=13,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=12
))

# Dòng header bảng
to_insert.append(build_para(
    "Từ viết tắt          Diễn giải tiếng Anh                     Ý nghĩa tiếng Việt",
    bold=True, size=13, underline=True,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_before=6, space_after=4
))

# Dòng kẻ ngang (dùng border bottom)
sep_p = build_para("", space_before=0, space_after=6)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
pPr_sep = sep_p.find(qn('w:pPr'))
if pPr_sep is None:
    pPr_sep = OxmlElement('w:pPr')
    sep_p.insert(0, pPr_sep)
pBdr = OxmlElement('w:pBdr')
bot  = OxmlElement('w:bottom')
bot.set(qn('w:val'),   'single')
bot.set(qn('w:sz'),    '6')
bot.set(qn('w:space'), '1')
bot.set(qn('w:color'), '1F4E79')
pBdr.append(bot)
pPr_sep.append(pBdr)
to_insert.append(sep_p)

# Các dòng từ viết tắt
for abbr, full_en, full_vi in ABBREVIATIONS:
    to_insert.append(build_abbr_row(abbr, full_en, full_vi))

# Dòng trống cuối
to_insert.append(build_para("", space_before=0, space_after=6))

# ---- Chèn tất cả vào trước MỞ ĐẦU ----
for elem in to_insert:
    mo_dau_elem.addprevious(elem)

# ---- Lưu ----
doc.save(out_path)
print(f"✅ Đã chèn {len(ABBREVIATIONS)} từ viết tắt")
print(f"✅ Lưu tại: {out_path}")

# ---- Xác minh ngay ----
doc2 = Document(out_path)
abbr_found = sum(1 for p in doc2.paragraphs if p.text.strip() and
                 any(a[0] in p.text for a in ABBREVIATIONS))
print(f"✅ Xác minh: tìm thấy {abbr_found} dòng có từ viết tắt")
print(f"✅ Tổng paragraphs trong file FINAL: {len(doc2.paragraphs)}")
