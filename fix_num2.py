# -*- coding: utf-8 -*-
"""Sửa nốt 2 số mục còn sai bằng cách quét toàn bộ runs"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

# Tìm file
for root, dirs, files in os.walk(r'D:\ViralWindow_Phan_Mem_Nhom_Kinh'):
    for f in files:
        if 'FINAL' in f and f.endswith('.docx') and not f.startswith('~'):
            path = os.path.join(root, f)

print("Path:", path)
doc = Document(path)
c = 0

for para in doc.paragraphs:
    full = para.text
    # Ktra toan bo doan van
    if '3.3.2.' in full and ('Trang' in full or 'trang' in full):
        for run in para.runs:
            if '3.3.2.' in run.text:
                run.text = run.text.replace('3.3.2.', '3.4.2.')
                c += 1
                print(f"Fixed run: {run.text[:60]}")
    if '3.4.3.' in full and ('Cài đặt' in full or 'Cai dat' in full or 'Cai' in full):
        for run in para.runs:
            if '3.4.3.' in run.text:
                run.text = run.text.replace('3.4.3.', '3.5.3.')
                c += 1
                print(f"Fixed run: {run.text[:60]}")

# Also check using para text merge approach for split runs
for para in doc.paragraphs:
    full_text = ''.join(r.text for r in para.runs)
    if '3.3.2' in full_text and 'Giao' in full_text and 'Trang' in full_text:
        # Rebuild numbering in first run containing '3.3.2'
        for run in para.runs:
            if '3.3' in run.text:
                old = run.text
                run.text = run.text.replace('3.3.2', '3.4.2')
                if old != run.text:
                    c += 1
                    print(f"Fixed merged: {run.text[:60]}")
    if '3.4.3' in full_text and 'Giao' in full_text and 'Cài' in full_text:
        for run in para.runs:
            if '3.4.3' in run.text:
                old = run.text
                run.text = run.text.replace('3.4.3', '3.5.3')
                if old != run.text:
                    c += 1
                    print(f"Fixed merged: {run.text[:60]}")

doc.save(path)
print(f"\nDone: {c} fixes")

# Verify
doc2 = Document(path)
print("\nChuong 3 structure:")
for p in doc2.paragraphs:
    t = p.text.strip()
    if len(t) < 70 and t and t[0] == '3' and '.' in t[:5]:
        print(f"  {t}")
