# -*- coding: utf-8 -*-
import sys
import os

# Try to use python-docx
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Try to use win32com as fallback
try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

doc_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main.docx"
out_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\thesis_extracted.txt"

if HAS_DOCX:
    print("Using python-docx")
    doc = Document(doc_path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)
    text = "\n".join(lines)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"SUCCESS: {len(text)} chars extracted")
    print(f"Paragraphs: {len(lines)}")

elif HAS_WIN32:
    print("Using win32com")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(doc_path)
    text = doc.Content.Text
    doc.Close(False)
    word.Quit()
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"SUCCESS: {len(text)} chars extracted")

else:
    # Try zip-based extraction (docx is a zip)
    print("Using zip extraction")
    import zipfile
    import xml.etree.ElementTree as ET
    
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    def get_text_from_xml(xml_content):
        root = ET.fromstring(xml_content)
        texts = []
        for para in root.iter(f'{ns}p'):
            para_text = ''
            for t in para.iter(f'{ns}t'):
                if t.text:
                    para_text += t.text
            if para_text.strip():
                texts.append(para_text)
        return "\n".join(texts)
    
    with zipfile.ZipFile(doc_path, 'r') as z:
        with z.open('word/document.xml') as f:
            xml_content = f.read()
    
    text = get_text_from_xml(xml_content)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"SUCCESS: {len(text)} chars extracted")
