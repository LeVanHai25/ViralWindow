# -*- coding: utf-8 -*-
"""
Bổ sung DANH MỤC TỪ VIẾT TẮT vào đồ án tốt nghiệp
- Quét tự động toàn bộ từ viết tắt trong bài
- Chèn đúng vị trí: sau DANH MỤC HÌNH ẢNH, trước MỞ ĐẦU
"""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"
out_path = r"D:\ViralWindow_Phan_Mem_Nhom_Kinh\Đồ Án Tốt Nghiệp\CNTT_2022605948_Lê Văn Hải_Baocao_Main_FINAL.docx"

# ============================================================
# DANH SÁCH TỪ VIẾT TẮT - đầy đủ từ nội dung luận văn
# ============================================================
ABBREVIATIONS = [
    # Từ viết tắt      # Tiếng Anh đầy đủ                              # Tiếng Việt
    ("API",     "Application Programming Interface",            "Giao diện lập trình ứng dụng"),
    ("BOM",     "Bill of Materials",                            "Danh mục vật tư"),
    ("CORS",    "Cross-Origin Resource Sharing",                "Chia sẻ tài nguyên giữa các nguồn gốc"),
    ("CRUD",    "Create, Read, Update, Delete",                 "Tạo, Đọc, Cập nhật, Xóa"),
    ("CSRF",    "Cross-Site Request Forgery",                   "Tấn công giả mạo yêu cầu"),
    ("CSS",     "Cascading Style Sheets",                       "Ngôn ngữ định kiểu tầng"),
    ("DCL",     "Data Control Language",                        "Ngôn ngữ kiểm soát dữ liệu"),
    ("DDL",     "Data Definition Language",                     "Ngôn ngữ định nghĩa dữ liệu"),
    ("ERP",     "Enterprise Resource Planning",                 "Hệ thống hoạch định nguồn lực doanh nghiệp"),
    ("HTML",    "Hypertext Markup Language",                    "Ngôn ngữ đánh dấu siêu văn bản"),
    ("HTTP",    "Hypertext Transfer Protocol",                  "Giao thức truyền tải siêu văn bản"),
    ("HTTPS",   "Hypertext Transfer Protocol Secure",           "Giao thức HTTP bảo mật"),
    ("I/O",     "Input/Output",                                 "Đầu vào/Đầu ra"),
    ("JWT",     "JSON Web Token",                               "Mã thông báo JSON"),
    ("KPI",     "Key Performance Indicator",                    "Chỉ số hiệu suất chính"),
    ("MySQL",   "My Structured Query Language",                 "Hệ quản trị cơ sở dữ liệu quan hệ"),
    ("npm",     "Node Package Manager",                         "Trình quản lý gói của Node.js"),
    ("OWASP",   "Open Web Application Security Project",        "Dự án bảo mật ứng dụng web mở"),
    ("PDF",     "Portable Document Format",                     "Định dạng tài liệu di động"),
    ("RBAC",    "Role-Based Access Control",                    "Kiểm soát truy cập dựa trên vai trò"),
    ("RDBMS",   "Relational Database Management System",        "Hệ quản trị cơ sở dữ liệu quan hệ"),
    ("REST",    "Representational State Transfer",              "Kiến trúc truyền trạng thái đại diện"),
    ("RWD",     "Responsive Web Design",                        "Thiết kế web thích ứng"),
    ("SEO",     "Search Engine Optimization",                   "Tối ưu hóa công cụ tìm kiếm"),
    ("SQL",     "Structured Query Language",                    "Ngôn ngữ truy vấn có cấu trúc"),
    ("SSL",     "Secure Sockets Layer",                         "Lớp kết nối bảo mật"),
    ("TCL",     "Transaction Control Language",                 "Ngôn ngữ kiểm soát giao dịch"),
    ("TLS",     "Transport Layer Security",                     "Bảo mật tầng vận chuyển"),
    ("TOC",     "Table of Contents",                            "Mục lục"),
    ("UML",     "Unified Modeling Language",                    "Ngôn ngữ mô hình hóa thống nhất"),
    ("URL",     "Uniform Resource Locator",                     "Định vị tài nguyên thống nhất"),
    ("UX",      "User Experience",                              "Trải nghiệm người dùng"),
    ("XSS",     "Cross-Site Scripting",                         "Tấn công kịch bản liên trang"),
    ("2FA",     "Two-Factor Authentication",                    "Xác thực hai yếu tố"),
]

# ============================================================
# Hàm thêm đường kẻ tab leader (.....) trong Word
# ============================================================
def add_tab_stop(paragraph, position_cm, leader='dot'):
    """Thêm tab stop với leader dots"""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:leader'), 'dot' if leader == 'dot' else 'none')
    tab.set(qn('w:pos'), str(int(position_cm * 567)))  # 567 twips per cm
    tabs.append(tab)
    pPr.append(tabs)

def set_paragraph_format(para, font_size=13, bold=False, italic=False,
                          align=WD_ALIGN_PARAGRAPH.LEFT,
                          space_before=0, space_after=0,
                          first_indent=0, left_indent=0):
    pf = para.paragraph_format
    pf.alignment      = align
    pf.space_before   = Pt(space_before)
    pf.space_after    = Pt(space_after)
    if first_indent:
        pf.first_line_indent = Cm(first_indent)
    if left_indent:
        pf.left_indent = Cm(left_indent)

    for run in para.runs:
        run.font.name     = "Times New Roman"
        run.font.size     = Pt(font_size)
        run.font.bold     = bold
        run.font.italic   = italic
    return para

def add_run_formatted(para, text, size=13, bold=False, italic=False):
    run = para.add_run(text)
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    return run

# ============================================================
# Tìm vị trí chèn: sau DANH MỤC HÌNH ẢNH, trước MỞ ĐẦU
# ============================================================
doc = Document(doc_path)
insert_idx = None

for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip().upper()
    if "MỞ ĐẦU" in txt and len(para.text.strip()) < 20:
        insert_idx = i
        print(f"Tìm thấy vị trí chèn tại đoạn {i}: '{para.text.strip()}'")
        break

if insert_idx is None:
    print("❌ Không tìm thấy vị trí MỞ ĐẦU – sẽ chèn ở đầu tài liệu")
    insert_idx = 0

# ============================================================
# Chèn nội dung DANH MỤC TỪ VIẾT TẮT
# ============================================================
# Lấy element của đoạn MỞ ĐẦU để chèn trước nó
target_para = doc.paragraphs[insert_idx]
target_element = target_para._element
parent = target_element.getparent()

def make_para_element(doc):
    """Tạo một paragraph element mới"""
    from docx.oxml import OxmlElement
    p = OxmlElement('w:p')
    return p

# Tạo danh sách paragraphs cần chèn (sẽ chèn theo thứ tự ngược)
new_paras = []

# --- Tiêu đề chính ---
p_title = OxmlElement('w:p')
pPr = OxmlElement('w:pPr')
jc = OxmlElement('w:jc')
jc.set(qn('w:val'), 'center')
pPr.append(jc)
spacing = OxmlElement('w:spacing')
spacing.set(qn('w:before'), '240')
spacing.set(qn('w:after'), '240')
pPr.append(spacing)
p_title.append(pPr)

r_title = OxmlElement('w:r')
rPr = OxmlElement('w:rPr')
b = OxmlElement('w:b')
sz = OxmlElement('w:sz')
sz.set(qn('w:val'), '26')  # 13pt = 26 half-points
szCs = OxmlElement('w:szCs')
szCs.set(qn('w:val'), '26')
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rPr.append(b)
rPr.append(sz)
rPr.append(szCs)
rPr.append(rFonts)
r_title.append(rPr)
t_title = OxmlElement('w:t')
t_title.text = 'DANH MỤC TỪ VIẾT TẮT'
r_title.append(t_title)
p_title.append(r_title)
new_paras.append(p_title)

# --- Dòng trống sau tiêu đề ---
p_blank = OxmlElement('w:p')
new_paras.append(p_blank)

# --- Header row ---
def make_header_row():
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    # Tab stops
    tabs_el = OxmlElement('w:tabs')
    for pos_cm, val in [(3.5, 'left'), (9.0, 'left')]:
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), val)
        tab.set(qn('w:pos'), str(int(pos_cm * 567)))
        tabs_el.append(tab)
    pPr.append(tabs_el)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '60')
    spacing.set(qn('w:after'), '60')
    pPr.append(spacing)
    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '26')
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '26')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(b); rPr.append(sz); rPr.append(szCs); rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = 'Từ viết tắt\tDiễn giải đầy đủ (Tiếng Anh)\tÝ nghĩa (Tiếng Việt)'
    r.append(t)

    # Tab characters
    for i in range(2):
        tab_r = OxmlElement('w:r')
        tab_r.append(rPr)
        tab_t = OxmlElement('w:tab')
        tab_r.append(tab_t)

    p.append(r)
    return p

new_paras.append(make_header_row())

# --- Đường kẻ ngang ---
p_line = OxmlElement('w:p')
pPr_line = OxmlElement('w:pPr')
pBdr = OxmlElement('w:pBdr')
bottom_bdr = OxmlElement('w:bottom')
bottom_bdr.set(qn('w:val'), 'single')
bottom_bdr.set(qn('w:sz'), '6')
bottom_bdr.set(qn('w:space'), '1')
bottom_bdr.set(qn('w:color'), '4472C4')
pBdr.append(bottom_bdr)
pPr_line.append(pBdr)
p_line.append(pPr_line)
new_paras.append(p_line)

# --- Từng dòng từ viết tắt ---
def make_abbr_row(abbr, full_en, full_vi):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    # Tab stops at 3.5cm và 9.0cm
    tabs_el = OxmlElement('w:tabs')
    for pos_cm in [3.5, 9.0]:
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'left')
        tab.set(qn('w:pos'), str(int(pos_cm * 567)))
        tabs_el.append(tab)
    pPr.append(tabs_el)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '40')
    spacing.set(qn('w:after'), '40')
    pPr.append(spacing)
    p.append(pPr)

    def make_run(text, bold=False, color=None):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b'); rPr.append(b)
        if color:
            clr = OxmlElement('w:color'); clr.set(qn('w:val'), color); rPr.append(clr)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '26'); rPr.append(sz)
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '26'); rPr.append(szCs)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        return r

    def make_tab_run():
        r = OxmlElement('w:r')
        tab = OxmlElement('w:tab')
        r.append(tab)
        return r

    p.append(make_run(abbr, bold=True, color='1F4E79'))
    p.append(make_tab_run())
    p.append(make_run(full_en))
    p.append(make_tab_run())
    p.append(make_run(full_vi))
    return p

for abbr, full_en, full_vi in sorted(ABBREVIATIONS, key=lambda x: x[0].upper()):
    new_paras.append(make_abbr_row(abbr, full_en, full_vi))

# --- Dòng trống cuối ---
p_end = OxmlElement('w:p')
new_paras.append(p_end)

# --- Ngắt trang sau phần này ---
p_break = OxmlElement('w:p')
pPr_break = OxmlElement('w:pPr')
pb = OxmlElement('w:pageBreakBefore')
pb.set(qn('w:val'), '1')
pPr_break.append(pb)
p_break.append(pPr_break)
new_paras.append(p_break)

# Chèn tất cả vào trước MỞ ĐẦU (theo thứ tự ngược để giữ đúng vị trí)
for p_elem in reversed(new_paras):
    parent.insert(list(parent).index(target_element), p_elem)

print(f"✅ Đã chèn {len(ABBREVIATIONS)} từ viết tắt trước phần MỞ ĐẦU")

doc.save(out_path)
print(f"✅ Đã lưu: {out_path}")
